"""Post-process a Pandoc-generated DOCX to add Chinese thesis formatting.

Pandoc's DOCX pipeline cannot produce:
- Different headers on different sections
- Mixed Roman/Arabic page numbering
- Table of contents fields
- Three-line table format (三线表)
- Chapter page breaks

This script uses python-docx to add these elements after Pandoc generation.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from html import escape
import re
import sys
from pathlib import Path
import zipfile
from lxml import etree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Pt

from scripts.chinese_fonts import resolve_cjk_font

_SONG = resolve_cjk_font("宋体")
_HEI = resolve_cjk_font("黑体")
_KAI = resolve_cjk_font("楷体")
_LISHU = resolve_cjk_font("隶书")


# Profile-specific settings — fallback only; primary source is PROFICES from export_chinese_thesis.
PROFILE_HEADER_TEXT: dict[str, str] = {
    "mainland-fallback": "本科毕业论文",
    "guangxi-undergrad": "广西大学本科毕业论文",
    "sichuan-grad": "四川大学硕士学位论文",
}

PROFILE_TOC_TITLE: dict[str, str] = {
    "mainland-fallback": "目　录",
    "guangxi-undergrad": "目　录",
    "sichuan-grad": "目　录",
}

# Margins in twips (1/20 point; 1440 twips = 1 inch = 2.54 cm)
PROFILE_MARGINS: dict[str, dict[str, int]] = {
    "mainland-fallback": {"top": 1440, "bottom": 1440, "left": 1440, "right": 1440, "header": 720, "footer": 720},
    # Guangxi: A4, top/bottom 2.54 cm, left/right 2.2 cm
    "guangxi-undergrad": {"top": 1440, "bottom": 1440, "left": 1247, "right": 1247, "header": 720, "footer": 720},
    # Sichuan: 16开 (184mm×260mm), top/bottom 2.5 cm, left 2.8 cm, right 2.2 cm
    "sichuan-grad":     {"top": 1417, "bottom": 1417, "left": 1587, "right": 1247, "header": 720, "footer": 720},
}

# Paper sizes in twips (1440 twips = 1 inch).
PROFILE_PAPER_SIZES: dict[str, dict[str, int]] = {
    "mainland-fallback": {"w": 11906, "h": 16838},  # A4
    "guangxi-undergrad": {"w": 11906, "h": 16838},  # A4
    "sichuan-grad":     {"w": 10430, "h": 14740},  # 16开 184mm × 260mm
}

# Profile-specific heading formats. All sizes in half-points (12pt = 24).
# See _format_headings_and_body() for the consumer.
PROFILE_HEADING_FORMATS: dict[str, dict[str, dict[str, str]]] = {
    "guangxi-undergrad": {
        "Heading1": {"sz": "36", "align": "center", "before": "400", "after": "400"},
        "Heading2": {"sz": "30", "align": "left",   "before": "0",   "after": "0"},
        "Heading3": {"sz": "28", "align": "left",   "before": "150",   "after": "0"},
        "Heading4": {"sz": "24", "align": "left",   "before": "0",   "after": "0"},
        "heading4_east_asia": _HEI,
    },
    "sichuan-grad": {
        # Sichan Level 1: 小三黑体(15pt), left-aligned
        "Heading1": {"sz": "30", "align": "left",  "before": "0",   "after": "0"},
        "Heading2": {"sz": "28", "align": "left",  "before": "0",   "after": "0"},
        "Heading3": {"sz": "24", "align": "left",  "before": "0",   "after": "0"},
        "Heading4": {"sz": "24", "align": "left",  "before": "0",   "after": "0"},
        "heading4_east_asia": _KAI,  # Sichuan Level 4 uses 楷体, not 黑体
    },
    "mainland-fallback": {
        "Heading1": {"sz": "36", "align": "center", "before": "400", "after": "400"},
        "Heading2": {"sz": "30", "align": "left",   "before": "0",   "after": "0"},
        "Heading3": {"sz": "28", "align": "left",   "before": "0",   "after": "0"},
        "Heading4": {"sz": "24", "align": "left",   "before": "0",   "after": "0"},
        "heading4_east_asia": _HEI,
    },
}

# Import PROFILES lazily so the module can be used standalone.
_PROFILES: dict | None = None


def _get_export_profiles():
    """Lazy import of PROFILES to avoid circular import at module load time."""
    global _PROFILES
    if _PROFILES is None:
        try:
            from scripts.export_chinese_thesis import PROFILES as _p
            _PROFILES = _p
        except ImportError:
            _PROFILES = {}
    return _PROFILES


def _profile_header_text(profile_id: str) -> str:
    """Get header text from the profile object, falling back to the static dict."""
    profiles = _get_export_profiles()
    p = profiles.get(profile_id)
    if p is not None:
        return p.header_text
    return PROFILE_HEADER_TEXT.get(profile_id, "本科毕业论文")


def _profile_toc_title(profile_id: str) -> str:
    """Get TOC title from the profile object, falling back to the static dict."""
    profiles = _get_export_profiles()
    p = profiles.get(profile_id)
    if p is not None:
        return p.toc_title
    return PROFILE_TOC_TITLE.get(profile_id, "目  录")

COVER_LABEL_TO_FIELD: dict[str, str] = {
    "学院": "college",
    "专业": "major",
    "班级": "class-name",
    "学号": "student-id",
    "姓名": "author",
    "指导老师": "advisor",
    "指导教师": "advisor",
    "日期": "date",
}


def _set_cell_font(run, font_name: str, font_size: Pt) -> None:
    """Set font on a run element."""
    run.font.name = font_name
    run.font.size = font_size
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _find_title(doc: Document) -> str:
    """Extract paper title from the document."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 4:
            return text[:60]
    return "论文题目"


def _set_cell_text(cell, text: str) -> None:
    """Replace a table cell's text while keeping the cell container."""
    cell.text = text


def _normalized_label(text: str) -> str:
    return re.sub(r"[\s:：]+", "", text)


_CHINESE_DIGITS = "〇一二三四五六七八九"


def _normalize_date_to_chinese(date: str) -> str:
    """Normalize a date string to the Chinese-numeral format used on thesis covers.

    Accepts both Chinese (``二〇二六年六月``) and Arabic (``2026年6月`` /
    ``2026-06`` / ``2026/6``) inputs and returns the Chinese form, e.g.
    ``二〇二六年六月``.  Strings that are already in Chinese form or that
    do not match a recognizable date pattern are returned unchanged.

    .. note::
       Only the year and month are preserved.  A trailing day component
       (e.g. ``2026年6月15日``) is intentionally dropped because Chinese
       thesis covers conventionally display only year + month.
    """
    if not date:
        return date

    text = date.strip()
    # Already Chinese-numeral form — return as-is.
    if re.match(r"^[二〇零一二三四五六七八九十]{4}年", text):
        return text

    # Arabic year + month: "2026年6月", "2026年06月", "2026-06", "2026/6"
    m = re.match(r"^(\d{4})\s*[年\-/年.]\s*(\d{1,2})\s*月?", text)
    if m:
        year = m.group(1)
        month = int(m.group(2))
        cn_year = "".join(_CHINESE_DIGITS[int(ch)] for ch in year)
        cn_month = _chinese_month(month)
        return f"{cn_year}年{cn_month}"

    # Year only: "2026年"
    m = re.match(r"^(\d{4})\s*年\s*$", text)
    if m:
        year = m.group(1)
        cn_year = "".join(_CHINESE_DIGITS[int(ch)] for ch in year)
        return f"{cn_year}年"

    return text


def _chinese_month(month: int) -> str:
    """Convert an Arabic month number to the Chinese month expression."""
    month_names = {
        1: "一月", 2: "二月", 3: "三月", 4: "四月", 5: "五月", 6: "六月",
        7: "七月", 8: "八月", 9: "九月", 10: "十月", 11: "十一月", 12: "十二月",
    }
    return month_names.get(month, f"{month}月")


def _make_title_paragraph_xml(title: str) -> str:
    """Build a standalone title paragraph: 黑体 一号(26pt), centered, underlined."""
    return (
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:line="480" w:lineRule="auto"/>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
        f'      <w:b/><w:sz w:val="52"/><w:szCs w:val="52"/>'
        f'      <w:u w:val="single"/>'
        f'    </w:rPr>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
        f'      <w:b/><w:sz w:val="52"/><w:szCs w:val="52"/>'
        f'      <w:u w:val="single"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{escape(title)}</w:t>'
        f'  </w:r>'
        f"</w:p>"
    )


def _extract_title_from_cover_table(cover_doc: Document, title: str) -> bool:
    """Remove the 课题名称 table and insert a standalone title paragraph.

    The cover template places 课题名称 inside a 2-row table with merged cells.
    Per university formatting norms, the title should appear as an independent
    paragraph *before* the field table (学院/专业/班级/...).

    Returns True if a title table was found and replaced.
    """
    if not title:
        return False
    body = cover_doc.element.body
    for table in list(cover_doc.tables):
        labels = {_normalized_label(cell.text) for row in table.rows for cell in row.cells}
        if "课题名称" not in labels:
            continue
        tbl_elem = table._tbl
        insert_idx = list(body).index(tbl_elem)
        body.remove(tbl_elem)
        body.insert(insert_idx, parse_xml(_make_title_paragraph_xml(title)))
        return True
    return False


def _fill_cover_doc(cover_doc: Document, cover_fields: dict[str, str]) -> None:
    """Fill known official cover fields without inventing missing values.

    Applies Chinese-university thesis formatting to filled fields:
    - Title: extracted from 课题名称 table → 黑体 一号(26pt) paragraph before field table
    - Advisor: 宋体 四号(14pt)
    - Other fields: inherit template formatting
    - Date: 宋体 四号(14pt), centered
    """
    title = cover_fields.get("title", "")
    _extract_title_from_cover_table(cover_doc, title)

    filled: set[str] = set()
    rows_to_delete: list = []
    normalized_date = _normalize_date_to_chinese(cover_fields.get("date", ""))
    for table in cover_doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = _normalized_label(row.cells[0].text)
            field = COVER_LABEL_TO_FIELD.get(label)
            if field and cover_fields.get(field):
                if field in filled:
                    rows_to_delete.append(row._tr)
                    continue
                value = normalized_date if field == "date" else cover_fields[field]
                _set_cell_text(row.cells[1], value)
                if field == "advisor":
                    for para in row.cells[1].paragraphs:
                        for run in para.runs:
                            _set_run_font(run, east_asia=_SONG, ascii="Times New Roman",
                                          sz="28", bold=False)
                elif field == "date":
                    for para in row.cells[1].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            _set_run_font(run, east_asia=_SONG, ascii="Times New Roman",
                                          sz="28", bold=False)
                filled.add(field)
    for tr in rows_to_delete:
        parent = tr.getparent()
        if parent is not None:
            parent.remove(tr)

    # Date fallback: if date was not filled via a table row, place it as a
    # standalone paragraph near the bottom of the cover (after the last
    # field table).  The date is normalized to Chinese numerals.
    date = normalized_date
    if date and "date" not in filled:
        _insert_date_paragraph(cover_doc, date)


def _insert_date_paragraph(cover_doc: Document, date: str) -> None:
    """Insert or replace the date paragraph at the bottom of the cover.

    The date should appear *after* the last field table (学院/专业/...),
    not in an arbitrary empty paragraph elsewhere on the cover.  We first
    try to find an existing date-matching or empty paragraph that sits
    after the last table; if none is found we insert a new centered date
    paragraph there.
    """
    date_pattern = re.compile(
        r"^[二〇零一二三四五六七八九十\d]{4}年"
        r"|^[\dX×]{4}\s*年"
        r"|^XXXX\s*年"
        r"|^\d{4}\s*年"
    )
    body = cover_doc.element.body
    children = list(body)

    # Find the index of the last w:tbl element in the body.
    last_tbl_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:tbl"):
            last_tbl_idx = i

    if last_tbl_idx is None:
        # No table at all — fall back to searching all paragraphs in reverse.
        search_start = len(children)
    else:
        search_start = last_tbl_idx + 1

    # Collect paragraphs after the last table (or all if no table).
    trailing_paras = [
        (i, child)
        for i, child in enumerate(children)
        if child.tag == qn("w:p") and i >= search_start
    ]

    # Try to find an existing date or empty paragraph after the last table.
    for idx, elem in reversed(trailing_paras):
        # Skip the very last paragraph if it holds the sectPr.
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            continue
        text = _paragraph_text(elem)
        if not text or date_pattern.match(text):
            _set_paragraph_date_text(elem, date)
            return

    # No suitable paragraph found — insert a new one right after the last
    # table (or before the last paragraph if there is no table).
    date_xml = _date_paragraph_xml(date)
    if last_tbl_idx is not None:
        children[last_tbl_idx].addnext(parse_xml(date_xml))
    elif children:
        # Insert before the last paragraph (usually the sectPr carrier).
        last_para = children[-1]
        if last_para.tag == qn("w:p"):
            last_para.addprevious(parse_xml(date_xml))
        else:
            body.append(parse_xml(date_xml))
    else:
        body.append(parse_xml(date_xml))


def _date_paragraph_xml(date: str) -> str:
    return (
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:line="360" w:lineRule="auto"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="{_SONG}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'      <w:sz w:val="28"/><w:szCs w:val="28"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{escape(date)}</w:t>'
        f'  </w:r>'
        f"</w:p>"
    )


def _set_paragraph_date_text(elem, date: str) -> None:
    """Replace the text of an existing paragraph with *date*, centered."""
    for t_elem in elem.findall(".//" + qn("w:t")):
        t_elem.text = ""
    t_list = elem.findall(".//" + qn("w:t"))
    if t_list:
        t_list[0].text = date
    else:
        run_xml = (
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr>'
            f'    <w:rFonts w:eastAsia="{_SONG}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'    <w:sz w:val="28"/><w:szCs w:val="28"/>'
            f'  </w:rPr>'
            f'  <w:t xml:space="preserve">{escape(date)}</w:t>'
            f"</w:r>"
        )
        elem.append(parse_xml(run_xml))
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        elem.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    else:
        jc.set(qn("w:val"), "center")


def _body_children_without_section(doc: Document) -> list:
    return [child for child in doc.element.body if child.tag != qn("w:sectPr")]


def _strip_section_properties(element) -> None:
    """Remove copied section references that point to another DOCX package."""
    for sect_pr in list(element.findall(".//" + qn("w:sectPr"))):
        parent = sect_pr.getparent()
        if parent is not None:
            parent.remove(sect_pr)


def _strip_comment_markers(element) -> None:
    """Remove comment markers from a cover element before inserting into the
    output document.  Cover template comments are editorial metadata that do
    not belong in the student's final thesis, and stripping them here
    eliminates any dangling-comment-marker risk."""
    for marker in list(element.findall(".//" + qn("w:commentRangeStart"))):
        parent = marker.getparent()
        if parent is not None:
            parent.remove(marker)
    for marker in list(element.findall(".//" + qn("w:commentRangeEnd"))):
        parent = marker.getparent()
        if parent is not None:
            parent.remove(marker)
    # w:commentReference sits inside w:rPr — remove the whole rPr if it only had this child
    for marker in list(element.findall(".//" + qn("w:commentReference"))):
        parent = marker.getparent()   # w:rPr
        gp = parent.getparent() if parent is not None else None
        parent.remove(marker)
        # If rPr is now empty, remove it too.
        if parent is not None and len(parent) == 0 and gp is not None:
            gp.remove(parent)


def _paragraph_has_content(elem) -> bool:
    """Check if a paragraph element contains any visible content.

    Returns True if the paragraph has text (w:t), a drawing (w:drawing),
    a picture (w:pict), an OLE object (w:object), or a math element (m:oMath).
    """
    if _paragraph_text(elem).strip():
        return True
    for tag in ("w:drawing", "w:pict", "w:object"):
        if elem.find(".//" + qn(tag)) is not None:
            return True
    if elem.find(".//" + qn("m:oMath")) is not None:
        return True
    return False


def _compact_cover_top(cover_doc: Document) -> None:
    """Remove redundant leading blank paragraphs from a cover template.

    Keeps at most one blank paragraph between any two consecutive *content*
    paragraphs so the cover fits on one page instead of spilling onto a
    second.  Content paragraphs are those with text or embedded drawings
    (e.g. the school logo) — see ``_paragraph_has_content``.

    The compaction runs across the whole cover body (up to the first table),
    collapsing runs of two or more blank paragraphs down to a single blank.
    Paragraphs that contain drawings (e.g. the school logo) are never
    removed (Issue #78).
    """
    body = cover_doc.element.body
    paragraphs = [c for c in body if c.tag == qn("w:p")]

    # Only compact paragraphs that appear before the first table; content
    # after tables (field values, date, section break) is left untouched.
    first_table_idx = None
    for i, child in enumerate(body):
        if child.tag == qn("w:tbl"):
            first_table_idx = i
            break
    if first_table_idx is not None:
        leading_paras = [p for p in paragraphs if body.index(p) < first_table_idx]
    else:
        leading_paras = paragraphs

    if len(leading_paras) <= 1:
        return

    # Walk through the leading paragraphs and collapse consecutive blanks.
    # Keep at most one blank between two content paragraphs.
    blanks_in_run: list = []
    to_remove: list = []
    for elem in leading_paras:
        if _paragraph_has_content(elem):
            # End of a blank run — collapse to at most one blank.
            if len(blanks_in_run) > 1:
                to_remove.extend(blanks_in_run[1:])
            blanks_in_run = []
        else:
            blanks_in_run.append(elem)

    # Trailing blanks at the very end of the leading region (just before the
    # first table) are also collapsed to one.
    if len(blanks_in_run) > 1:
        to_remove.extend(blanks_in_run[1:])

    for elem in to_remove:
        body.remove(elem)


def _insert_cover(
    doc: Document,
    cover_docx: Path,
    cover_fields: dict[str, str],
    margins: dict[str, int] | None = None,
    paper_size: dict[str, int] | None = None,
) -> int:
    """Prepend a filled DOCX cover template, then insert a section break.

    The section break (not just a page break) creates a new section for
    the abstract pages that follow, allowing independent page numbering.
    """
    if margins is None:
        margins = {"top": 1440, "bottom": 1440, "left": 1440, "right": 1440,
                    "header": 720, "footer": 720}
    if paper_size is None:
        paper_size = {"w": 11906, "h": 16838}  # A4
    cover_doc = Document(str(cover_docx))
    _fill_cover_doc(cover_doc, cover_fields)
    _compact_cover_top(cover_doc)

    body = doc.element.body
    insert_at = 0
    for element in _body_children_without_section(cover_doc):
        copied = deepcopy(element)
        _strip_section_properties(copied)
        _strip_comment_markers(copied)
        body.insert(insert_at, copied)
        insert_at += 1

    # Section break between cover and abstract — starts a new section with
    # profile-specific margins and Roman page numbering (hidden on cover,
    # visible on abstract pages).
    sect_break = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:sectPr>'
        f'      <w:pgSz w:w="{paper_size["w"]}" w:h="{paper_size["h"]}"/>'
        f'      <w:pgMar w:top="{margins["top"]}" w:bottom="{margins["bottom"]}"'
        f'              w:left="{margins["left"]}" w:right="{margins["right"]}"'
        f'              w:header="{margins["header"]}" w:footer="{margins["footer"]}"/>'
        f'    </w:sectPr>'
        f'  </w:pPr>'
        f"</w:p>"
    )
    body.insert(insert_at, sect_break)
    return insert_at + 1


def _merge_cover_package_resources(output_docx: Path, cover_docx: Path) -> None:
    """Merge media files, relationships, and content types from cover DOCX into output.

    python-docx does not propagate image relationships or media when raw XML
    elements are inserted from another document.  This function opens the saved
    DOCX as a ZIP archive and merges the cover's package parts so that images
    (e.g. school logos) render correctly.
    """
    # 1. Read cover package resources.
    with zipfile.ZipFile(cover_docx, "r") as cz:
        cover_rels_raw = cz.read("word/_rels/document.xml.rels")
        cover_media = {name: cz.read(name) for name in cz.namelist() if name.startswith("word/media/")}

    # 2. Parse cover relationships to find image rels.
    NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
    cover_rels = etree.fromstring(cover_rels_raw)
    cover_image_rels: list[tuple[str, str]] = []
    IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    for child in cover_rels:
        rel_type = child.get("Type", "")
        target = child.get("Target", "")
        if rel_type == IMAGE_REL_TYPE:
            cover_image_rels.append((child.get("Id", ""), target))

    if not cover_image_rels:
        return

    # 3. Read output package: max existing rId, existing media, document.xml.
    with zipfile.ZipFile(output_docx, "r") as oz:
        output_namelist = oz.namelist()
        output_rels_raw = oz.read("word/_rels/document.xml.rels")
        output_ct_raw = oz.read("[Content_Types].xml")
        output_doc_raw = oz.read("word/document.xml")

    # 4. Parse output relationships to find max rId.
    output_rels = etree.fromstring(output_rels_raw)
    max_rId = 0
    for child in output_rels:
        rid = child.get("Id", "")
        m = re.match(r"rId(\d+)", rid)
        if m:
            max_rId = max(max_rId, int(m.group(1)))

    existing_media = {name for name in output_namelist if name.startswith("word/media/")}

    # 5. Build rId remap from cover rels to new output rels.
    rId_remap: dict[str, str] = {}

    for old_rid, target in cover_image_rels:
        target_path = f"word/{target}" if not target.startswith("word/") else target
        if target_path in existing_media:
            # Already exists; reuse existing relationship.
            for child in output_rels:
                if child.get("Target", "") == target:
                    rId_remap[old_rid] = child.get("Id", "")
                    break
            continue

        max_rId += 1
        new_rid = f"rId{max_rId}"
        rId_remap[old_rid] = new_rid
        rel_elem = etree.SubElement(output_rels, "Relationship")
        rel_elem.set("Id", new_rid)
        rel_elem.set("Type", IMAGE_REL_TYPE)
        rel_elem.set("Target", target)

    if not rId_remap:
        return

    # 6. Replace old cover rIds in document.xml body with remapped ones.
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    output_doc = etree.fromstring(output_doc_raw)
    body = output_doc.find(f"{{{NS_W}}}body")
    if body is not None:
        for elem in body.iter():
            for attr_name in (f"{{{NS_R}}}embed", f"{{{NS_R}}}id"):
                old_val = elem.get(attr_name)
                if old_val and old_val in rId_remap:
                    elem.set(attr_name, rId_remap[old_val])

    output_doc_bytes = etree.tostring(
        output_doc, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # 7. Merge content types.
    CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
    output_ct = etree.fromstring(output_ct_raw)
    known_parts = {child.get("PartName", "") for child in output_ct}

    CONTENT_TYPE_MAP = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tiff": "image/tiff",
        "tif": "image/tiff",
        "svg": "image/svg+xml",
        "wmf": "image/x-wmf",
        "emf": "image/x-emf",
    }
    for name in sorted(cover_media):
        part_name = f"/{name}"
        if part_name in known_parts:
            continue
        ext = name.rsplit(".", 1)[-1].lower()
        ct = CONTENT_TYPE_MAP.get(ext, "application/octet-stream")
        override = etree.SubElement(output_ct, f"{{{CT_NS}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", ct)

    output_ct_bytes = etree.tostring(
        output_ct, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # 8. Serialize updated relationships.
    # Remove default namespace declaration from rels root (OOXML requires unqualified attrs).
    output_rels_bytes = etree.tostring(
        output_rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    _replace_in_zip(
        output_docx,
        {
            "word/document.xml": output_doc_bytes,
            "word/_rels/document.xml.rels": output_rels_bytes,
            "[Content_Types].xml": output_ct_bytes,
            **cover_media,
        },
    )


def _replace_in_zip(zip_path: Path, replacements: dict[str, bytes]) -> None:
    """Replace or add files in an existing ZIP archive."""
    tmp_path = zip_path.with_suffix(zip_path.suffix + ".tmp")
    with zipfile.ZipFile(zip_path, "r") as src:
        existing_names = set(src.namelist())
        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                name = item.filename
                if name in replacements:
                    dst.writestr(name, replacements[name])
                else:
                    dst.writestr(name, src.read(name))
            for name, data in replacements.items():
                if name not in existing_names:
                    dst.writestr(name, data)
    tmp_path.replace(zip_path)


def _run_properties_xml(
    *,
    bold: bool = False,
    size: int | None = None,
    east_asia_font: str = _SONG,
    ascii_font: str = "Times New Roman",
) -> str:
    rpr_parts = [
        f'<w:rFonts w:eastAsia="{east_asia_font}" w:ascii="{ascii_font}" w:hAnsi="{ascii_font}"/>'
    ]
    if bold:
        rpr_parts.append("<w:b/>")
    if size:
        rpr_parts.append(f'<w:sz w:val="{size}"/>')
        rpr_parts.append(f'<w:szCs w:val="{size}"/>')
    return f"<w:rPr>{''.join(rpr_parts)}</w:rPr>"


def _paragraph_xml(
    text: str,
    *,
    align: str | None = None,
    bold: bool = False,
    size: int | None = None,
    east_asia_font: str = _SONG,
    ascii_font: str = "Times New Roman",
    before: int | None = None,
    after: int | None = None,
    first_line_chars: int | None = None,
    outline_lvl: int | None = None,
) -> str:
    ppr_parts: list[str] = []
    spacing_attrs = ['w:line="400"', 'w:lineRule="exact"']
    if before is not None:
        spacing_attrs.append(f'w:before="{before}"')
    if after is not None:
        spacing_attrs.append(f'w:after="{after}"')
    ppr_parts.append(f"<w:spacing {' '.join(spacing_attrs)}/>")
    if first_line_chars is not None:
        ppr_parts.append(f'<w:ind w:firstLineChars="{first_line_chars}"/>')
    if align:
        ppr_parts.append(f'<w:jc w:val="{align}"/>')
    if outline_lvl is not None:
        ppr_parts.append(f'<w:outlineLvl w:val="{outline_lvl}"/>')
    rpr = _run_properties_xml(
        bold=bold,
        size=size,
        east_asia_font=east_asia_font,
        ascii_font=ascii_font,
    )
    ppr = f"<w:pPr>{''.join(ppr_parts)}{rpr}</w:pPr>" if ppr_parts or rpr else ""
    return f'<w:p {nsdecls("w")}>{ppr}<w:r>{rpr}<w:t>{escape(text)}</w:t></w:r></w:p>'


def _blank_line_xml() -> str:
    return (
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:line="400" w:lineRule="exact"/></w:pPr>'
        f"</w:p>"
    )


def _strip_trailing_punctuation(text: str) -> str:
    return text.strip().rstrip("；;，,。.")


def _capitalize_english_keywords(text: str) -> str:
    terms = []
    for term in re.split(r"\s*;\s*", _strip_trailing_punctuation(text)):
        term = term.strip()
        if not term:
            continue
        terms.append(term[:1].upper() + term[1:])
    return "; ".join(terms)


def _capitalize_english_title(text: str) -> str:
    words = []
    for word in text.split():
        words.append(word if word.isupper() else word[:1].upper() + word[1:])
    return " ".join(words)


def _keyword_paragraph_xml(
    label: str,
    terms: str,
    *,
    label_east_asia_font: str,
    term_east_asia_font: str,
    ascii_font: str = "Times New Roman",
    first_line_chars: int | None = None,
) -> str:
    ppr_parts = ['<w:spacing w:line="400" w:lineRule="exact"/>']
    if first_line_chars is not None:
        ppr_parts.append(f'<w:ind w:firstLineChars="{first_line_chars}"/>')
    label_rpr = _run_properties_xml(
        bold=True,
        size=24,
        east_asia_font=label_east_asia_font,
        ascii_font=ascii_font,
    )
    term_rpr = _run_properties_xml(
        bold=False,
        size=24,
        east_asia_font=term_east_asia_font,
        ascii_font=ascii_font,
    )
    return (
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>{"".join(ppr_parts)}{term_rpr}</w:pPr>'
        f'  <w:r>{label_rpr}<w:t>{escape(label)}</w:t></w:r>'
        f'  <w:r>{term_rpr}<w:t>{escape(terms)}</w:t></w:r>'
        f"</w:p>"
    )


def _page_break_xml() -> str:
    return f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>'


def _insert_xml_elements(doc: Document, insert_at: int, xml_chunks: list[str]) -> int:
    body = doc.element.body
    for chunk in xml_chunks:
        body.insert(insert_at, parse_xml(chunk))
        insert_at += 1
    return insert_at


def _split_abstract_paragraphs(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n\s*\n", text.strip()) if part.strip()]


def _insert_abstract_pages(doc: Document, insert_at: int, abstract_fields: dict[str, str]) -> int:
    """Insert Chinese and English abstract pages after the cover."""
    title_en = abstract_fields.get("title-en", "").strip()
    abstract_zh = abstract_fields.get("abstract-zh", "").strip()
    keywords_zh = _strip_trailing_punctuation(abstract_fields.get("keywords-zh", ""))
    abstract_en = abstract_fields.get("abstract-en", "").strip()
    keywords_en = _capitalize_english_keywords(abstract_fields.get("keywords-en", ""))

    chunks: list[str] = [
        _blank_line_xml(),
        _paragraph_xml("摘　要", align="center", bold=True, size=32, east_asia_font=_HEI, outline_lvl=1),
        _blank_line_xml(),
    ]
    chunks.extend(
        _paragraph_xml(paragraph, first_line_chars=200)
        for paragraph in _split_abstract_paragraphs(abstract_zh)
    )
    chunks.append(_blank_line_xml())
    chunks.append(
        _keyword_paragraph_xml(
            "关键词：",
            keywords_zh,
            label_east_asia_font=_HEI,
            term_east_asia_font=_SONG,
            first_line_chars=200,
        )
    )
    chunks.append(_page_break_xml())
    if title_en:
        chunks.append(
            _paragraph_xml(
                _capitalize_english_title(title_en),
                align="center",
                bold=True,
                size=32,
                east_asia_font="Times New Roman",
                ascii_font="Times New Roman",
            )
        )
        chunks.append(_blank_line_xml())
    chunks.append(
        _paragraph_xml(
            "ABSTRACT",
            align="center",
            bold=True,
            size=32,
            east_asia_font="Times New Roman",
            ascii_font="Times New Roman",
            outline_lvl=1,
        )
    )
    chunks.append(_blank_line_xml())
    chunks.extend(
        _paragraph_xml(
            paragraph,
            first_line_chars=400,
            east_asia_font="Times New Roman",
            ascii_font="Times New Roman",
        )
        for paragraph in _split_abstract_paragraphs(abstract_en)
    )
    chunks.append(_blank_line_xml())
    chunks.append(
        _keyword_paragraph_xml(
            "Keywords: ",
            keywords_en,
            label_east_asia_font="Times New Roman",
            term_east_asia_font="Times New Roman",
            ascii_font="Times New Roman",
        )
    )
    chunks.append(_page_break_xml())
    return _insert_xml_elements(doc, insert_at, chunks)


def _is_cover_metadata_table(table) -> bool:
    labels = {_normalized_label(cell.text) for row in table.rows for cell in row.cells}
    return bool({"课题名称", "学院", "专业", "班级", "学号", "姓名", "指导老师"} & labels)


def _paragraph_text(element) -> str:
    texts = element.findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


def _remove_generated_title_paragraph(doc: Document, title: str) -> None:
    """Remove Pandoc's generated metadata title paragraph from the body."""
    if not title:
        return
    body = doc.element.body
    for elem in list(body):
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem)
        if not text:
            continue
        if text == title:
            body.remove(elem)
            return


def _add_header_to_section(section, header_text: str, title_text: str) -> None:
    """Add header with university name (left) and title (right), 隶书小四号."""
    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_left = para.add_run(header_text)
    _set_cell_font(run_left, _LISHU, Pt(12))

    tab_stops = para.paragraph_format.tab_stops
    if tab_stops:
        try:
            tab_stops.add_tab_stop(Cm(14.7), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        except Exception:
            pass

    run_tab = para.add_run("\t")
    run_right = para.add_run(title_text)
    _set_cell_font(run_right, _LISHU, Pt(12))

    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def _add_page_number_footer(section, is_roman: bool = True) -> None:
    """Add page number footer: Times New Roman, centered."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    _set_cell_font(run, "Times New Roman", Pt(10.5))
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._element.append(fld_char_begin)

    run2 = para.add_run()
    _set_cell_font(run2, "Times New Roman", Pt(10.5))
    instr = ""
    if is_roman:
        instr = 'PAGE  \\* ROMAN'
    else:
        instr = "PAGE"
    instr_text = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">{instr}</w:instrText>'
    )
    run2._element.append(instr_text)

    run3 = para.add_run()
    _set_cell_font(run3, "Times New Roman", Pt(10.5))
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run3._element.append(fld_char_end)


def _is_figure_table(table) -> bool:
    """Check if a table is a Pandoc figure-fallback table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return False
    tblStyle = tblPr.find(qn("w:tblStyle"))
    if tblStyle is None:
        return False
    return tblStyle.get(qn("w:val")) == "FigureTable"


def _format_tables_three_line(doc: Document) -> None:
    """Format all tables as 三线表 (three-line table).

    三线表 rules:
    - Top border: thick (1.5 pt, sz=24)
    - Bottom border: thick (1.5 pt, sz=24)
    - Header row bottom border: thin (0.5 pt, sz=6)
    - No left, right, inside vertical, or inside horizontal borders
    - Header row: bold
    """
    for table in doc.tables:
        if _is_cover_metadata_table(table) or _is_figure_table(table):
            continue
        if not table.rows:
            continue

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            tbl.insert(0, tblPr)

        # Remove existing borders element if any
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        # Build three-line borders
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="24" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="24" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(parse_xml(borders_xml))

        # Set table width to 100% of page
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn("w:w"), "5000")
            tblW.set(qn("w:type"), "pct")

        # Add bottom border to header row (first row) — thin line
        if len(table.rows) > 1:
            first_row = table.rows[0]
            for cell in first_row.cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                    tc.insert(0, tcPr)

                # Remove existing cell borders
                existing = tcPr.find(qn("w:tcBorders"))
                if existing is not None:
                    tcPr.remove(existing)

                cell_borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(cell_borders)

            # Bold header row
            for cell in first_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _format_table_caption_paragraph(caption_elem) -> None:
    """Apply Guangxi University formatting to a table caption paragraph."""
    cap_pPr = caption_elem.find(qn("w:pPr"))
    if cap_pPr is None:
        cap_pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        caption_elem.insert(0, cap_pPr)

    # Center alignment
    if cap_pPr.find(qn("w:jc")) is None:
        cap_pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    # Fixed 20 pt line spacing (400 twips)
    spacing = cap_pPr.find(qn("w:spacing"))
    if spacing is None:
        cap_pPr.append(
            parse_xml(f'<w:spacing {nsdecls("w")} w:line="400" w:lineRule="exact"/>')
        )
    else:
        spacing.set(qn("w:line"), "400")
        spacing.set(qn("w:lineRule"), "exact")

    # Bold + 五号 (10.5pt = 21 half-pts) on all runs
    for run in caption_elem.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            run.insert(0, rPr)
        if rPr.find(qn("w:b")) is None:
            rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        sz = rPr.find(qn("w:sz"))
        if sz is None:
            rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="21"/>'))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="21"/>'))


def _format_table_cell_content(table) -> None:
    """Set table cell font, line spacing, and remove first-line indent."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pPr = para._element.find(qn("w:pPr"))
                if pPr is not None:
                    # Remove first-line indent
                    ind = pPr.find(qn("w:ind"))
                    if ind is not None:
                        ind.set(qn("w:firstLineChars"), "0")
                        ind.set(qn("w:firstLine"), "0")
                    # Fixed 20 pt line spacing
                    spacing = pPr.find(qn("w:spacing"))
                    if spacing is None:
                        pPr.append(
                            parse_xml(f'<w:spacing {nsdecls("w")} w:line="400" w:lineRule="exact"/>')
                        )
                    else:
                        spacing.set(qn("w:line"), "400")
                        spacing.set(qn("w:lineRule"), "exact")

                # Set font on all runs
                for run in para.runs:
                    _set_cell_font(run, _SONG, Pt(10.5))


def _format_table_captions_and_content(doc: Document) -> None:
    """Format table captions and cell content per Guangxi University requirements.

    - Caption above table, centered, 五号 (10.5pt) bold, fixed 20 pt line spacing
    - One blank line before caption and after table
    - Table content: 五号 (10.5pt) 宋体, no first-line indent
    """
    body = doc.element.body
    children = list(body)

    # Collect (caption_idx, table_idx) pairs for captioned tables
    pairs: list[tuple[int, int]] = []
    i = 0
    while i < len(children):
        elem = children[i]
        if elem.tag == qn("w:p"):
            pPr = elem.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None and pStyle.get(qn("w:val")) == "TableCaption":
                    if i + 1 < len(children) and children[i + 1].tag == qn("w:tbl"):
                        pairs.append((i, i + 1))
                        i += 2
                        continue
        i += 1

    # Format cell content for all non-excluded tables
    for table in doc.tables:
        if _is_cover_metadata_table(table):
            continue
        _format_table_cell_content(table)

    if not pairs:
        return

    # Process captions in reverse order
    blank = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:line="400" w:lineRule="exact"/></w:pPr>'
        f"</w:p>"
    )
    for cap_idx, tbl_idx in reversed(pairs):
        caption_elem = children[cap_idx]
        _format_table_caption_paragraph(caption_elem)

        # Insert blank before caption
        body.insert(cap_idx, deepcopy(blank))
        # Insert blank after table (tbl_idx+1 because cap_idx insert shifted tbl_idx by 1)
        body.insert(tbl_idx + 1, deepcopy(blank))


def _format_figure_captions(doc: Document) -> None:
    """Format figure captions per university requirements.

    - Caption below the figure, centered, 五号 (10.5pt) bold, fixed 20 pt line spacing
    - One blank line before the figure and after the caption

    Pandoc generates figures as CaptionedFigure + ImageCaption paragraph pairs.
    """
    body = doc.element.body
    children = list(body)

    pairs: list[tuple[int, int]] = []
    i = 0
    while i < len(children):
        elem = children[i]
        if elem.tag != qn("w:p"):
            i += 1
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "CaptionedFigure":
                if i + 1 < len(children):
                    next_elem = children[i + 1]
                    if next_elem.tag == qn("w:p"):
                        next_pPr = next_elem.find(qn("w:pPr"))
                        if next_pPr is not None:
                            next_pStyle = next_pPr.find(qn("w:pStyle"))
                            if next_pStyle is not None and next_pStyle.get(qn("w:val")) == "ImageCaption":
                                pairs.append((i, i + 1))
                                i += 2
                                continue
        i += 1

    if not pairs:
        return

    for fig_idx, cap_idx in reversed(pairs):
        caption_elem = children[cap_idx]
        cap_pPr = caption_elem.find(qn("w:pPr"))
        if cap_pPr is None:
            cap_pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            caption_elem.insert(0, cap_pPr)
        if cap_pPr.find(qn("w:jc")) is None:
            cap_pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
        spacing = cap_pPr.find(qn("w:spacing"))
        if spacing is None:
            cap_pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:line="400" w:lineRule="exact"/>'))
        else:
            spacing.set(qn("w:line"), "400")
            spacing.set(qn("w:lineRule"), "exact")
        for run in caption_elem.findall(qn("w:r")):
            rPr = run.find(qn("w:rPr"))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                run.insert(0, rPr)
            if rPr.find(qn("w:b")) is None:
                rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="21"/>')
                rPr.append(sz)
            szCs = rPr.find(qn("w:szCs"))
            if szCs is None:
                szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="21"/>')
                rPr.append(szCs)
        blank = (
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr><w:spacing w:line="400" w:lineRule="exact"/></w:pPr>'
            f"</w:p>"
        )
        body.insert(cap_idx + 1, parse_xml(blank))
        body.insert(fig_idx, parse_xml(blank))


def _get_paragraph_all_text(elem) -> str:
    """Extract all text from a paragraph including OMML math text (m:t)."""
    parts: list[str] = []
    for child in elem.iter():
        local = child.tag.rsplit("}", 1)[-1]  # strip namespace
        if local == "t" and child.text:
            parts.append(child.text)
    return "".join(parts)


def _format_equation_numbers(doc: Document) -> None:
    """Add chapter-scoped equation numbers at the right margin.

    Equations before the first chapter heading use global numbers:
      (1), (2), ...

    Equations after a chapter heading use chapter-scoped numbers with
    the chapter counter reset at each heading:
      (1-1), (1-2), …, (2-1), (2-2), …

    Equations with an existing \\tag{...} or a matching (N-N) label are
    skipped (preserving authored tags).  Appendix chapters use letter
    prefixes: (A-1), (B-1), …
    """
    body = doc.element.body

    # Read page dimensions from the first sectPr
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        return

    pgSz = sectPr.find(qn("w:pgSz"))
    pgMar = sectPr.find(qn("w:pgMar"))

    page_width = 11906  # A4 default in twips
    left_margin = 1701  # ~3cm default
    right_margin = 1701

    if pgSz is not None:
        w = pgSz.get(qn("w:w"))
        if w is not None:
            page_width = int(w)
    if pgMar is not None:
        left = pgMar.get(qn("w:left"))
        right = pgMar.get(qn("w:right"))
        if left is not None:
            left_margin = int(left)
        if right is not None:
            right_margin = int(right)

    right_tab_pos = page_width - right_margin
    if right_tab_pos <= left_margin:
        right_tab_pos = 9638  # sensible fallback

    chapter_counter = 0
    equation_counter = 0
    appendix_counter = 0

    for elem in body:
        if elem.tag != qn("w:p"):
            continue

        # Detect chapter heading boundaries to reset equation counter
        texts = elem.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts).strip()
        if _is_chapter_heading(text):
            equation_counter = 0
            normalized = re.sub(r"\s+", "", text)
            if normalized.startswith("附录"):
                appendix_counter += 1
            elif re.match(r"^第[一二三四五六七八九十\d]+章", normalized):
                chapter_counter += 1
            # 参考文献/致谢: counter reset only, no chapter increment
            continue

        omath_para = elem.find(qn("m:oMathPara"))
        if omath_para is None:
            continue

        # Skip equations that already carry an authored tag
        all_text = _get_paragraph_all_text(elem)
        if re.search(r'\\tag\s*\{', all_text):
            continue
        if re.search(r'\([\dA-Z]+-\d+\)', all_text):
            continue

        equation_counter += 1

        # Generate chapter-scoped or global number
        if appendix_counter > 0:
            prefix = chr(ord("A") + appendix_counter - 1)
            num_text = f"({prefix}-{equation_counter})"
        elif chapter_counter > 0:
            num_text = f"({chapter_counter}-{equation_counter})"
        else:
            num_text = f"({equation_counter})"

        # Add right-aligned tab stop in paragraph properties
        pPr = elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            elem.insert(0, pPr)

        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = parse_xml(f'<w:tabs {nsdecls("w")}></w:tabs>')
            pPr.append(tabs)
        tabs.append(parse_xml(
            f'<w:tab {nsdecls("w")} w:val="right" w:pos="{right_tab_pos}"/>'
        ))

        # Ensure fixed 20pt line spacing
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            pPr.append(parse_xml(
                f'<w:spacing {nsdecls("w")} w:line="400" w:lineRule="exact"/>'
            ))
        else:
            spacing.set(qn("w:line"), "400")
            spacing.set(qn("w:lineRule"), "exact")

        # Insert tab character + equation number after oMathPara
        run_xml = (
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr>'
            f'    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'    <w:sz w:val="21"/>'
            f'    <w:szCs w:val="21"/>'
            f'  </w:rPr>'
            f'  <w:tab/>'
            f'  <w:t xml:space="preserve">{num_text}</w:t>'
            f'</w:r>'
        )
        omath_para.addnext(parse_xml(run_xml))


def _is_toc_field_paragraph(elem) -> bool:
    """Check if a paragraph element contains a TOC field (fldChar + TOC instr)."""
    for instr in elem.findall(".//" + qn("w:instrText")):
        if instr.text and "TOC" in instr.text:
            return True
    return False


def _remove_existing_toc_entries(doc: Document) -> None:
    """Remove static/hand-written TOC entries and old TOC fields from the body.

    Pandoc or a reference template may pre-generate static TOC paragraphs
    (styled as 'toc 1', 'toc 2', 'toc 3', or 'TOCHeading') or even a stale
    TOC field.  These coexist badly with the auto TOC field inserted by
    ``_add_toc``, producing duplicate or non-updating entries.

    This function scans the document body and removes:
    - Paragraphs whose style is ``toc 1``, ``toc 2``, ``toc 3`` or
      ``TOCHeading``.
    - Paragraphs that contain a TOC field (``fldChar`` + ``instrText`` with
      ``TOC``).

    It does **not** remove the ``目  录`` heading itself — ``_add_toc``
    reuses or creates that heading.  The ``Compact`` style (a Pandoc
    body-text variant) is intentionally *not* removed because it may be
    used for legitimate non-TOC content.
    """
    body = doc.element.body
    TOC_ENTRY_STYLES = {"toc 1", "toc 2", "toc 3", "TOCHeading"}
    to_remove: list = []
    for elem in body:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) in TOC_ENTRY_STYLES:
                to_remove.append(elem)
                continue
        if _is_toc_field_paragraph(elem):
            to_remove.append(elem)
    for elem in to_remove:
        body.remove(elem)


def _add_toc(doc: Document, toc_title: str, insert_at: int | None = None) -> None:
    """Add a TOC field after the heading that matches the TOC title.

    When *insert_at* is provided, the TOC heading + field + page break are
    inserted at that child index in the document body.  Otherwise the TOC is
    placed before the front-matter-to-body section break (legacy path).

    TOC field shows only Level 1 and Level 2 (backslash o "1-2").
    """
    body_elem = doc.element.body
    children = list(body_elem)

    # Find existing TOC heading in document
    toc_heading = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.sub(r"\s+", "", text) == re.sub(r"\s+", "", toc_title):
            toc_heading = para
            break

    if toc_heading is not None:
        # Existing heading — insert TOC field after it, ensure break after
        toc_element = toc_heading._element
        parent = toc_element.getparent()
        idx = list(parent).index(toc_element)
    elif insert_at is not None:
        # Explicit insertion point (caller guarantees correct position)
        # Blank line before
        body_elem.insert(insert_at, parse_xml(_blank_line_xml()))
        insert_at += 1

        # TOC heading: 黑体 三号(32pt), centered
        heading_xml = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr>'
            f'    <w:jc w:val="center"/>'
            f'    <w:spacing w:before="400" w:after="400" w:line="400" w:lineRule="exact"/>'
            f'    <w:rPr><w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
            f'      <w:b/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr>'
            f'  </w:pPr>'
            f'  <w:r><w:rPr><w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
            f'      <w:b/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr>'
            f'    <w:t>{toc_title}</w:t></w:r>'
            f"</w:p>"
        )
        body_elem.insert(insert_at, heading_xml)
        toc_element = heading_xml
        parent = body_elem
        idx = insert_at
    else:
        # Legacy path: search for section break before first chapter
        first_chap_idx = None
        for i, elem in enumerate(children):
            if elem.tag != qn("w:p") or _is_toc_paragraph(elem):
                continue
            texts = elem.findall(".//" + qn("w:t"))
            text = "".join(t.text or "" for t in texts).strip()
            if _is_chapter_heading(text):
                first_chap_idx = i
                break

        sect_break_idx = None
        if first_chap_idx is not None:
            for i in range(first_chap_idx - 1, -1, -1):
                elem = children[i]
                if elem.tag != qn("w:p"):
                    continue
                pPr = elem.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    sect_break_idx = i
                    break

        insert_at_legacy = sect_break_idx if sect_break_idx is not None else (first_chap_idx or 0)

        # Blank line before
        body_elem.insert(insert_at_legacy, parse_xml(_blank_line_xml()))
        insert_at_legacy += 1

        # TOC heading: 黑体 三号(32pt), centered
        heading_xml = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr>'
            f'    <w:jc w:val="center"/>'
            f'    <w:spacing w:before="400" w:after="400" w:line="400" w:lineRule="exact"/>'
            f'    <w:rPr><w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
            f'      <w:b/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr>'
            f'  </w:pPr>'
            f'  <w:r><w:rPr><w:rFonts w:eastAsia="{_HEI}" w:ascii="{_HEI}" w:hAnsi="{_HEI}"/>'
            f'      <w:b/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr>'
            f'    <w:t>{toc_title}</w:t></w:r>'
            f"</w:p>"
        )
        body_elem.insert(insert_at_legacy, heading_xml)
        toc_element = heading_xml
        parent = body_elem
        idx = insert_at_legacy

    # TOC field paragraph (level 1-2 only)
    toc_xml = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        f'  <w:r><w:instrText xml:space="preserve">'
        f"    TOC \\o \"1-2\" \\h \\z \\u "
        f"  </w:instrText></w:r>"
        f'  <w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        f'  <w:r><w:t>[请在Word中右键更新目录]</w:t></w:r>'
        f'  <w:r><w:fldChar w:fldCharType="end"/></w:r>'
        f"</w:p>"
    )
    parent.insert(idx + 1, toc_xml)

    # Add page break after the TOC field to ensure TOC occupies its own page(s)
    # and body text starts on a fresh page (Issue #83)
    page_break_xml = parse_xml(_page_break_xml())
    parent.insert(idx + 2, page_break_xml)


def _format_toc_entries(doc: Document) -> None:
    """Add 'toc 1' and 'toc 2' styles so Word generates correctly-formatted TOC entries.

    When the user updates the TOC field in Word, entries use these styles.
    宋体 小四(12pt), left-aligned, no indent, fixed 20pt line spacing.
    """
    TOC_STYLE_XML = {
        "toc1": (
            f'<w:style {nsdecls("w")} w:type="paragraph" w:styleId="toc1">'
            f'  <w:name w:val="toc 1"/>'
            f'  <w:basedOn w:val="Normal"/>'
            f'  <w:pPr>'
            f'    <w:jc w:val="left"/>'
            f'    <w:spacing w:line="400" w:lineRule="exact"/>'
            f'  </w:pPr>'
            f'  <w:rPr>'
            f'    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
            f'               w:eastAsia="{_SONG}"/>'
            f'    <w:sz w:val="24"/><w:szCs w:val="24"/>'
            f'  </w:rPr>'
            f'</w:style>'
        ),
        "toc2": (
            f'<w:style {nsdecls("w")} w:type="paragraph" w:styleId="toc2">'
            f'  <w:name w:val="toc 2"/>'
            f'  <w:basedOn w:val="Normal"/>'
            f'  <w:pPr>'
            f'    <w:jc w:val="left"/>'
            f'    <w:spacing w:line="400" w:lineRule="exact"/>'
            f'  </w:pPr>'
            f'  <w:rPr>'
            f'    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
            f'               w:eastAsia="{_SONG}"/>'
            f'    <w:sz w:val="24"/><w:szCs w:val="24"/>'
            f'  </w:rPr>'
            f'</w:style>'
        ),
    }
    styles_elem = doc.styles.element
    existing_ids = {s.get(qn("w:styleId")) for s in styles_elem}
    for sid, xml in TOC_STYLE_XML.items():
        if sid not in existing_ids:
            styles_elem.append(parse_xml(xml))

    # Also format any existing TOC entry paragraphs (if Pandoc pre-generated them)
    TOC_STYLES = {"toc 1", "toc 2", "toc 3"}
    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        style = pStyle.get(qn("w:val"))
        if style not in TOC_STYLES:
            continue
        _set_para_alignment(pPr, "left")
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)
        for run in para.runs:
            _set_run_font(run, east_asia=_SONG, ascii="Times New Roman",
                          sz="24", bold=False)
        _set_para_line_spacing_fixed_20pt(pPr)


def _format_bibliography_entries(doc: Document) -> None:
    """Format Bibliography entries: 宋体 小四(12pt), fixed line spacing."""
    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        if pStyle.get(qn("w:val")) != "Bibliography":
            continue
        for run in para.runs:
            _set_run_font(run, east_asia=_SONG, ascii="Times New Roman",
                          sz="24", bold=False)
        _set_para_line_spacing_fixed_20pt(pPr)


def _reorder_bibliography_before_thanks(doc: Document) -> None:
    """Move Bibliography paragraphs after 致谢 to before 致谢.

    Pandoc citeproc places the reference list at the document end by default.
    This function moves Bibliography-styled paragraphs to before the 致谢
    heading (which should follow 参考文献 in the input markdown).
    """
    body = doc.element.body
    children = list(body)

    refs_heading_idx = None
    thanks_heading_idx = None
    for i, elem in enumerate(children):
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        style = pStyle.get(qn("w:val"))
        text = "".join(
            t.text or "" for t in elem.findall(".//" + qn("w:t"))
        ).strip()
        if style == "Heading1":
            if re.sub(r"\s+", "", text) == "参考文献":
                refs_heading_idx = i
            elif re.sub(r"\s+", "", text) == "致谢":
                thanks_heading_idx = i

    if refs_heading_idx is None or thanks_heading_idx is None:
        return  # nothing to reorder

    # Collect Bibliography paragraphs that are AFTER 致谢
    bib_elems = []
    for elem in children[thanks_heading_idx:]:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        if pStyle.get(qn("w:val")) == "Bibliography":
            bib_elems.append(elem)

    if not bib_elems:
        return

    # Move each Bibliography element to just before 致谢 heading
    for elem in bib_elems:
        body.remove(elem)
        body.insert(children.index(children[thanks_heading_idx]), elem)


def _format_headings_and_body(doc: Document, profile: str = "mainland-fallback") -> None:
    """Format headings and body paragraphs per profile-specific heading hierarchy.

    See PROFILE_HEADING_FORMATS for per-profile heading sizes and alignments.
    """
    heading_fmt = PROFILE_HEADING_FORMATS.get(profile, PROFILE_HEADING_FORMATS["mainland-fallback"])
    heading4_ea = heading_fmt.get("heading4_east_asia", _HEI)
    outline_levels = {"Heading1": "0", "Heading2": "1", "Heading3": "2", "Heading4": "3"}
    BODY_STYLES = {"Normal", "BodyText", "FirstParagraph", "First Paragraph"}
    EXCLUDED_BODY = {"Bibliography", "toc 1", "toc 2", "toc 3",
                     "Compact", "TableCaption", "ImageCaption",
                     "CaptionedFigure", "toc"}

    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle_elem = pPr.find(qn("w:pStyle"))
        style = pStyle_elem.get(qn("w:val")) if pStyle_elem is not None else ""

        if style in heading_fmt:
            # Special headings (参考文献/致谢) use 黑体 三号(16pt)
            clean_text = re.sub(r"\s+", "", para.text.strip())
            if clean_text in ("参考文献", "致谢"):
                _set_para_spacing(pPr, before="400", after="400")
                _set_para_alignment(pPr, "center")
                _set_para_line_spacing_fixed_20pt(pPr)
                _set_outline_level(pPr, "0")
                for run in para.runs:
                    _set_run_font(run, east_asia=_HEI, ascii=_HEI,
                                  sz="32", bold=True)
                continue
            fmt = heading_fmt[style]
            # --- paragraph properties ---
            _set_para_spacing(pPr, before=fmt["before"], after=fmt["after"])
            _set_para_alignment(pPr, fmt["align"])
            _set_para_line_spacing_fixed_20pt(pPr)
            _set_outline_level(pPr, outline_levels.get(style, "0"))
            # --- run properties (Level 4 may use 楷体 per profile spec) ---
            heading_font = heading4_ea if style == "Heading4" else _HEI
            for run in para.runs:
                _set_run_font(run, east_asia=heading_font, ascii=_HEI,
                              sz=fmt["sz"], bold=True)
        elif style in BODY_STYLES:
            if style in EXCLUDED_BODY:
                continue
            _set_para_alignment(pPr, "both")
            _set_para_line_spacing_fixed_20pt(pPr)
            _set_first_line_indent(pPr, chars=2)
            for run in para.runs:
                _set_run_font(run, east_asia=_SONG, ascii="Times New Roman",
                              sz="24", bold=False)


def _set_para_spacing(pPr, before: str, after: str) -> None:
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(
            f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}"/>'
        )
        pPr.append(spacing)
    else:
        spacing.set(qn("w:before"), before)
        spacing.set(qn("w:after"), after)


def _set_para_alignment(pPr, align: str) -> None:
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{align}"/>')
        pPr.append(jc)
    else:
        jc.set(qn("w:val"), align)


def _set_para_line_spacing_fixed_20pt(pPr) -> None:
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(
            f'<w:spacing {nsdecls("w")} w:line="400" w:lineRule="exact"/>'
        )
        pPr.append(spacing)
    else:
        spacing.set(qn("w:line"), "400")
        spacing.set(qn("w:lineRule"), "exact")


def _set_outline_level(pPr, level: str) -> None:
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{level}"/>')
        pPr.append(outline)
    else:
        outline.set(qn("w:val"), level)


def _set_style_line_spacing_exact(style, twips: str, rule: str = "exact") -> None:
    element = style.element
    pPr = element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        element.append(pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}></w:spacing>')
        pPr.append(spacing)
    spacing.set(qn("w:line"), twips)
    spacing.set(qn("w:lineRule"), rule)


def _set_style_east_asia_font(style, font_name: str) -> None:
    element = style.element
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_first_line_indent(pPr, chars: int = 2) -> None:
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:firstLineChars="{chars * 100}"/>'
        )
        pPr.append(ind)
    else:
        ind.set(qn("w:firstLineChars"), str(chars * 100))


def _set_run_font(run, east_asia: str, ascii: str, sz: str, bold: bool) -> None:
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'  w:eastAsia="{east_asia}" w:ascii="{ascii}" w:hAnsi="{ascii}"/>'
        )
        rPr.append(rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), east_asia)
        rFonts.set(qn("w:ascii"), ascii)
        rFonts.set(qn("w:hAnsi"), ascii)
    sz_elem = rPr.find(qn("w:sz"))
    if sz_elem is None:
        sz_elem = parse_xml(f'<w:sz {nsdecls("w")} w:val="{sz}"/>')
        rPr.append(sz_elem)
    else:
        sz_elem.set(qn("w:val"), sz)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{sz}"/>')
        rPr.append(szCs)
    else:
        szCs.set(qn("w:val"), sz)
    b_elem = rPr.find(qn("w:b"))
    if bold:
        if b_elem is None:
            rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
    else:
        if b_elem is not None:
            rPr.remove(b_elem)


def _is_chapter_heading(text: str) -> bool:
    """Check if paragraph text is a chapter-level heading.

    Matches patterns like:
    - 第1章 / 第N章 / 第一章
    - 参考文献, 致谢, 附录, 附录A

    Whitespace is normalized before matching so that spacing variants
    like "致  谢" are also recognized.
    """
    normalized = re.sub(r"\s+", "", text)
    if re.match(r"^第[一二三四五六七八九十\d]+章", normalized):
        return True
    if normalized.startswith("参考文献"):
        return True
    if normalized.startswith("致谢"):
        return True
    if normalized.startswith("附录"):
        return True
    return False


def _is_toc_paragraph(elem) -> bool:
    """Check if a paragraph element is a TOC entry (not a real heading).

    TOC entries have styles like 'toc 1', 'toc 2', 'toc 3' or 'Compact'.
    """
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    style_val = pStyle.get(qn("w:val"), "")
    if style_val.startswith("toc") or style_val == "Compact":
        return True
    return False


def _set_section_page_number_format(sect_pr, fmt: str) -> None:
    """Set page number format on a section's sectPr element."""
    pgNumType = sect_pr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        sect_pr.remove(pgNumType)
    pgNumType = parse_xml(
        f'<w:pgNumType {nsdecls("w")} w:fmt="{fmt}"/>'
    )
    sect_pr.append(pgNumType)


def _set_section_page_start(sect_pr, start: int) -> None:
    """Set page number starting value on a section's sectPr element."""
    pgNumType = sect_pr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = parse_xml(
            f'<w:pgNumType {nsdecls("w")} w:start="{start}"/>'
        )
        sect_pr.append(pgNumType)
    else:
        pgNumType.set(qn("w:start"), str(start))


def _add_section_breaks_and_page_breaks(doc: Document, margins: dict[str, int] | None = None, paper_size: dict[str, int] | None = None) -> None:
    """Add section breaks and page breaks for proper chapter separation.

    Creates a section break before the first chapter heading to separate
    front matter (Roman page numbers, no header) from body (Arabic page
    numbers, header). Adds pageBreakBefore to each subsequent chapter heading.
    Only actual headings are considered — TOC entries are ignored.

    margins: page margins dict with keys top/bottom/left/right/header/footer
             in twips. Defaults to 2.54 cm all sides, 0.5 inch header/footer.
    """
    if margins is None:
        margins = {"top": 1440, "bottom": 1440, "left": 1440, "right": 1440,
                    "header": 720, "footer": 720}
    if paper_size is None:
        paper_size = {"w": 11906, "h": 16838}  # A4
    body_elem = doc.element.body
    paragraphs = list(body_elem)

    first_chapter_idx = None
    first_chapter_elem = None

    # Find the first ACTUAL chapter heading (skip TOC entries)
    for i, elem in enumerate(paragraphs):
        if elem.tag != qn("w:p"):
            continue
        if _is_toc_paragraph(elem):
            continue
        texts = elem.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts).strip()
        if _is_chapter_heading(text):
            first_chapter_idx = i
            first_chapter_elem = elem
            break

    if first_chapter_elem is None:
        # No chapter heading found — add section break before last sectPr
        last_sect = body_elem.find(qn("w:sectPr"))
        if last_sect is not None:
            sect_para = parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'  <w:pPr><w:sectPr/></w:pPr>'
                f"</w:p>"
            )
            body_elem.insert(list(body_elem).index(last_sect), sect_para)
        return

    # Ensure the first chapter has a section break before it
    prev_elem = paragraphs[first_chapter_idx - 1] if first_chapter_idx > 0 else None

    has_sect_before = False
    if prev_elem is not None and prev_elem.tag == qn("w:p"):
        pPr = prev_elem.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            has_sect_before = True

    if not has_sect_before:
        sect_para = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr>'
            f'    <w:sectPr>'
            f'      <w:pgSz w:w="{paper_size["w"]}" w:h="{paper_size["h"]}"/>'
            f'      <w:pgMar w:top="{margins["top"]}" w:bottom="{margins["bottom"]}"'
            f'              w:left="{margins["left"]}" w:right="{margins["right"]}"'
            f'              w:header="{margins["header"]}" w:footer="{margins["footer"]}"/>'
            f'    </w:sectPr>'
            f'  </w:pPr>'
            f"</w:p>"
        )
        body_elem.insert(list(body_elem).index(first_chapter_elem), sect_para)

    # Add pageBreakBefore to subsequent chapters (skip TOC entries)
    chapter_count = 0
    for elem in body_elem:
        if elem.tag != qn("w:p"):
            continue
        if _is_toc_paragraph(elem):
            continue
        texts = elem.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts).strip()
        if _is_chapter_heading(text):
            chapter_count += 1
            if chapter_count > 1:
                pPr = elem.find(qn("w:pPr"))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    elem.insert(0, pPr)
                if pPr.find(qn("w:pageBreakBefore")) is None:
                    pPr.append(
                        parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>')
                    )


def _set_section_headers_footers(doc: Document, header_text: str, title: str) -> None:
    """Set headers and footers on each section.

    With cover (3 sections):
      Section 0: cover — no page numbers, no header
      Section 1: abstract — Roman page numbers (i, ii…), no header
      Section 2+: body — Arabic page numbers (1, 2…), header

    Without cover (2 sections):
      Section 0: front matter — Roman page numbers, no header
      Section 1: body — Arabic page numbers starting 1, header
    """
    sections = doc.sections
    n = len(sections)
    first_body_idx = 2 if n >= 3 else 1
    for i, section in enumerate(sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        if n >= 3 and i == 0:
            # Cover: no page numbers, no header
            for p in section.header.paragraphs:
                p.clear()
            for p in section.footer.paragraphs:
                p.clear()
        elif n >= 3 and i == 1:
            # Abstract: Roman page numbers, no header
            for p in section.header.paragraphs:
                p.clear()
            _add_page_number_footer(section, is_roman=True)
            _set_section_page_number_format(section._sectPr, "upperRoman")
            _set_section_page_start(section._sectPr, 1)
        elif i == 0:
            # Front matter (no cover): Roman page numbers, no header
            for p in section.header.paragraphs:
                p.clear()
            for p in section.footer.paragraphs:
                p.clear()
            _set_section_page_number_format(section._sectPr, "upperRoman")
        else:
            # Body: header + Arabic page numbers
            _add_header_to_section(section, header_text, title)
            _add_page_number_footer(section, is_roman=False)
            _set_section_page_number_format(section._sectPr, "decimal")
            if i == first_body_idx:
                pgNumType = section._sectPr.find(qn("w:pgNumType"))
                if pgNumType is not None:
                    pgNumType.set(qn("w:start"), "1")


def _auto_number_headings(doc: Document) -> None:
    """Add chapter and multi-level heading numbering to headings.

    Heading1: 第一章, 第二章, ... (Chinese numeral; skip 参考文献/致谢/附录)
    Heading2: 1.1, 1.2, ...
    Heading3: 1.1.1, 1.2.1, ...
    Heading4: 1.1.1.1, ...

    Strips any existing "第N章" or "N.N" prefix before adding fresh numbering.
    """
    CHAPTER_NUMERALS = [
        "零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
        "十一", "十二", "十三", "十四", "十五",
    ]
    EXISTING_CHAP = re.compile(r"^第[一二三四五六七八九十\d]+章\s*")
    EXISTING_SEC = re.compile(r"^(\d+\.)*\d+\s+")
    NO_NUMBER = {"参考文献", "致谢"}

    chap_counter = 0
    # [section, subsection, subsubsection] counters within current chapter
    sec = [0, 0, 0]

    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle_elem = pPr.find(qn("w:pStyle"))
        style = pStyle_elem.get(qn("w:val")) if pStyle_elem is not None else ""
        text = para.text.strip()
        if not text:
            continue

        if style == "Heading1":
            clean = re.sub(r"\s+", "", text)
            clean = EXISTING_CHAP.sub("", clean)
            if clean in NO_NUMBER or clean.startswith("附录"):
                continue
            chap_counter += 1
            sec = [0, 0, 0]
            stripped = EXISTING_CHAP.sub("", text).strip()
            new_text = f"第{CHAPTER_NUMERALS[chap_counter]}章 {stripped}"
            _replace_para_text(para, new_text)

        elif style == "Heading2" and chap_counter > 0:
            sec[0] += 1
            sec[1] = 0
            sec[2] = 0
            stripped = EXISTING_SEC.sub("", text).strip()
            new_text = f"{chap_counter}.{sec[0]} {stripped}"
            _replace_para_text(para, new_text)

        elif style == "Heading3" and chap_counter > 0:
            sec[1] += 1
            sec[2] = 0
            stripped = EXISTING_SEC.sub("", text).strip()
            new_text = f"{chap_counter}.{sec[0]}.{sec[1]} {stripped}"
            _replace_para_text(para, new_text)

        elif style == "Heading4" and chap_counter > 0:
            sec[2] += 1
            stripped = EXISTING_SEC.sub("", text).strip()
            new_text = f"{chap_counter}.{sec[0]}.{sec[1]}.{sec[2]} {stripped}"
            _replace_para_text(para, new_text)


def _replace_para_text(para, new_text: str) -> None:
    """Replace paragraph text, preserving existing runs for formatting."""
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
    else:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run_xml = (
            f'<w:r {nsdecls("w")}>{rPr}'
            f'<w:t xml:space="preserve">{escape(new_text)}</w:t></w:r>'
        )
        para._element.append(parse_xml(run_xml))


# --- Format Validation -------------------------------------------------------

# Per-profile validation expectations.
# Font names are the canonical Chinese names; comparison normalises
# platform variants (Songti SC → 宋体, Heiti SC → 黑体, Kaiti SC → 楷体,
# SimSun → 宋体, SimHei → 黑体, KaiTi → 楷体, STLiti → 隶书,
# Noto Serif/Sans CJK SC → 宋体/黑体).
_FONT_ALIASES: dict[str, str] = {
    "songti sc": "宋体", "simsun": "宋体", "noto serif cjk sc": "宋体",
    "heiti sc": "黑体", "simhei": "黑体", "noto sans cjk sc": "黑体",
    "kaiti sc": "楷体", "kaiti": "楷体",
    "times new roman": "times new roman",
    "隶书": "隶书", "lishu": "隶书", "stliti": "隶书",
}

def _normalise_font(name: str) -> str:
    return _FONT_ALIASES.get(name.strip().lower(), name.strip().lower())

PROFILE_FORMAT_SPECS: dict[str, dict] = {
    "guangxi-undergrad": {
        "body_font": "宋体",
        "body_font_ascii": "times new roman",
        "body_font_size_pt": 12,       # 小四
        "body_line_spacing_pt": 20,     # fixed
        "body_first_indent_chars": 2,
        "heading1_font": "黑体",
        "heading1_size_pt": 18,         # 小二号
        "heading2_font": "黑体",
        "heading2_size_pt": 15,         # 小三号
        "heading3_font": "黑体",
        "heading3_size_pt": 14,         # 四号
    },
    "sichuan-grad": {
        "body_font": "宋体",
        "body_font_ascii": "times new roman",
        "body_font_size_pt": 12,
        "body_line_spacing_pt": 20,
        "body_first_indent_chars": 2,
        "heading1_font": "黑体",
        "heading1_size_pt": 15,         # 小三号
        "heading2_font": "黑体",
        "heading2_size_pt": 14,         # 四号
        "heading3_font": "黑体",
        "heading3_size_pt": 12,         # 小四号
    },
    "mainland-fallback": {
        "body_font": "宋体",
        "body_font_ascii": "times new roman",
        "body_font_size_pt": 12,
        "body_line_spacing_pt": 20,
        "body_first_indent_chars": 2,
        "heading1_font": "黑体",
        "heading1_size_pt": 18,
        "heading2_font": "黑体",
        "heading2_size_pt": 15,
        "heading3_font": "黑体",
        "heading3_size_pt": 14,
    },
}


def validate_format(docx_path: Path, profile: str) -> list[dict]:
    """Validate a generated DOCX against the profile's format specifications.

    Returns a list of dicts, each describing one check result:
      - check: human-readable description of what was checked
      - expected: expected value per the profile specification
      - observed: actual value found in the document
      - result: "pass" | "fail" | "warn"
    """
    results: list[dict] = []
    spec = PROFILE_FORMAT_SPECS.get(profile)
    if spec is None:
        return [{"check": "profile validation", "expected": profile, "observed": "unknown profile", "result": "fail"}]

    expected_margins = PROFILE_MARGINS.get(profile, {})
    doc = Document(str(docx_path))

    # -- 1. Page margins (from document-level sectPr in the XML) --------------
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            for edge in ("top", "bottom", "left", "right"):
                exp = expected_margins.get(edge)
                if exp is None:
                    continue
                val_str = pgMar.get(qn(f"w:{edge}"))
                obs = int(val_str) if val_str else None
                if obs is not None:
                    # Convert twips → mm for readability
                    exp_mm = round(exp * 25.4 / 1440, 1)
                    obs_mm = round(obs * 25.4 / 1440, 1)
                    margin_ok = abs(obs - exp) <= 20  # ~1 pt tolerance
                    results.append({
                        "check": f"Page {edge} margin",
                        "expected": f"{exp_mm} mm ({exp} twips)",
                        "observed": f"{obs_mm} mm ({obs} twips)",
                        "result": "pass" if margin_ok else "fail",
                    })
                else:
                    results.append({
                        "check": f"Page {edge} margin",
                        "expected": f"{exp} twips",
                        "observed": "not found",
                        "result": "fail",
                    })
        else:
            results.append({"check": "Page margins", "expected": "present", "observed": "missing", "result": "fail"})
    else:
        results.append({"check": "Document section", "expected": "present", "observed": "missing", "result": "fail"})

    # -- 2. Normal (body) style: font, size, line spacing ---------------------
    normal_style = doc.styles["Normal"]
    _validate_style_font(results, normal_style, "body", spec)
    _validate_style_spacing(results, normal_style, spec)

    # -- 3. Heading styles ----------------------------------------------------
    for level, key_prefix in [(1, "heading1"), (2, "heading2"), (3, "heading3")]:
        style_id = f"Heading {level}"
        h_font = spec.get(f"{key_prefix}_font")
        h_size = spec.get(f"{key_prefix}_size_pt")
        if h_font is None or h_size is None:
            continue
        try:
            h_style = doc.styles[style_id]
        except KeyError:
            results.append({
                "check": f"Heading {level} style ({style_id})",
                "expected": "present",
                "observed": "missing",
                "result": "warn",
            })
            continue
        _validate_style_font(results, h_style, f"heading {level}", spec, key_prefix)

    # -- 4. Sections and page numbering ---------------------------------------
    # Count paragraph-embedded section breaks (each adds a new section)
    para_sectPrs = [
        p.find(qn("w:pPr")).find(qn("w:sectPr"))
        for p in body.findall(qn("w:p"))
        if p.find(qn("w:pPr")) is not None and p.find(qn("w:pPr")).find(qn("w:sectPr")) is not None
    ]
    section_count = 1 + len(para_sectPrs)
    results.append({
        "check": "Section count",
        "expected": ">= 3 (cover, abstract, body)",
        "observed": str(section_count),
        "result": "pass" if section_count >= 3 else "warn",
    })

    # Check page numbering on body sections
    pg_num_types = []
    body_sectPr = sectPr
    if body_sectPr is not None:
        pgNumType = body_sectPr.find(qn("w:pgNumType"))
        pg_num_types.append(pgNumType.get(qn("w:fmt")) if pgNumType is not None else "decimal(default)")
        pg_start = pgNumType.get(qn("w:start")) if pgNumType is not None else None
        results.append({
            "check": "Body page numbering format",
            "expected": "decimal (1, 2, 3...)",
            "observed": pg_num_types[-1],
            "result": "pass" if pg_num_types[-1] in ("decimal", "decimal(default)") else "warn",
        })

    results.append({
        "check": "Body page numbering start",
        "expected": "1",
        "observed": str(pg_start) if pg_start else "not set (default)",
        "result": "pass" if pg_start is None or pg_start == "1" else "warn",
    })

    # -- 5. Body paragraph sample (line spacing override at paragraph level) --
    _validate_body_paragraph_sample(results, doc, spec)

    return results


def _validate_body_paragraph_sample(results: list[dict], doc, spec: dict) -> None:
    """Spot-check one body paragraph for line-spacing override."""
    exp_spacing = spec.get("body_line_spacing_pt")
    if exp_spacing is None:
        return
    # Look for a paragraph AFTER the first chapter heading; all body content
    # follows chapter headings in Chinese thesis structure.
    skip_prefixes = ("abstract", "keywords", "keyword", "摘要", "关键词",
                     "目", "ABSTRACT", "KEYWORDS", "[", "TOC", "toc")
    found_chapter = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not found_chapter:
            if re.match(r"^第[一二三四五六七八九十\d]+章", text.replace(" ", "")):
                found_chapter = True
            continue
        if any(text.startswith(p) for p in skip_prefixes):
            continue
        if len(text) < 10:  # skip very short lines (sub-headings, labels)
            continue
        # Found a plausible body paragraph — check its spacing
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is not None:
                line_val = spacing.get(qn("w:line"))
                line_rule = spacing.get(qn("w:lineRule"))
                if line_rule == "exact" and line_val:
                    obs_pt = round(int(line_val) / 20, 1)
                    spacing_ok = abs(obs_pt - exp_spacing) < 2.0
                    results.append({
                        "check": "Body paragraph line spacing (sample)",
                        "expected": f"{exp_spacing} pt exact",
                        "observed": f"{obs_pt} pt exact",
                        "result": "pass" if spacing_ok else "fail",
                    })
                else:
                    obs_line = round(int(line_val) / 20, 1) if line_val else "not set"
                    results.append({
                        "check": "Body paragraph line spacing (sample)",
                        "expected": f"{exp_spacing} pt exact",
                        "observed": f"{obs_line} pt (rule={line_rule})",
                        "result": "fail",
                    })
            else:
                results.append({
                    "check": "Body paragraph line spacing (sample)",
                    "expected": f"{exp_spacing} pt exact",
                    "observed": "no spacing element (inherited from style)",
                    "result": "warn",
                })
        break  # Only check first body paragraph


def _validate_style_font(results: list[dict], style, label: str, spec: dict, prefix: str = "body") -> None:
    """Validate font name and size of a python-docx style against spec."""
    f = style.font
    exp_font = spec.get(f"{prefix}_font")
    exp_ascii = spec.get(f"{prefix}_font_ascii") if prefix == "body" else None
    exp_size = spec.get(f"{prefix}_font_size_pt")

    # Font size (half-points in DOCX; 12pt = 24 half-pts)
    if exp_size is not None:
        obs_sz = f.size
        if obs_sz is not None:
            obs_pt = obs_sz / 12700  # EMU → pt
            size_ok = abs(obs_pt - exp_size) < 1.0
            results.append({
                "check": f"{label.capitalize()} font size",
                "expected": f"{exp_size} pt",
                "observed": f"{obs_pt:.0f} pt",
                "result": "pass" if size_ok else "fail",
            })
        else:
            results.append({
                "check": f"{label.capitalize()} font size",
                "expected": f"{exp_size} pt",
                "observed": "not set (inherited)",
                "result": "pass",  # inheritance is normal
            })

    # East-Asian font
    if exp_font is not None:
        # Read via XML for east-asian font
        rPr = style.element.find(qn("w:rPr"))
        obs_ea = "not set"
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                ea_val = rFonts.get(qn("w:eastAsia"))
                if ea_val:
                    obs_ea = _normalise_font(ea_val)
        exp_normalised = _normalise_font(exp_font)
        results.append({
            "check": f"{label.capitalize()} CJK font",
            "expected": f"{exp_font} ({exp_normalised})",
            "observed": obs_ea,
            "result": "pass" if obs_ea == exp_normalised else "warn",
        })

    # ASCII font
    if exp_ascii is not None:
        rPr = style.element.find(qn("w:rPr"))
        obs_ascii = "not set"
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                ascii_val = rFonts.get(qn("w:ascii"))
                if ascii_val:
                    obs_ascii = ascii_val.strip().lower()
        results.append({
            "check": f"{label.capitalize()} ASCII font",
            "expected": exp_ascii,
            "observed": obs_ascii,
            "result": "pass" if obs_ascii == exp_ascii else "warn",
        })


def _validate_style_spacing(results: list[dict], style, spec: dict) -> None:
    """Validate line spacing of a python-docx style against spec."""
    exp_spacing = spec.get("body_line_spacing_pt")
    if exp_spacing is None:
        return
    pf = style.paragraph_format
    # line_spacing returns a float (lines) or None; line_spacing_rule is an enum
    ls = pf.line_spacing
    lsr = pf.line_spacing_rule
    if ls is not None:
        if lsr in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
            obs_pt = round(int(ls) / 12700, 1)
            spacing_ok = abs(obs_pt - exp_spacing) < 2.0
            results.append({
                "check": "Body line spacing (Normal style)",
                "expected": f"{exp_spacing} pt exact",
                "observed": f"{obs_pt} pt (rule={lsr})",
                "result": "pass" if spacing_ok else "warn",
            })
        else:
            results.append({
                "check": "Body line spacing (Normal style)",
                "expected": f"{exp_spacing} pt exact",
                "observed": f"{float(ls)} lines (rule={lsr})",
                "result": "warn",
            })
    else:
        results.append({
            "check": "Body line spacing",
            "expected": f"{exp_spacing} pt exact",
            "observed": "not set (inherited)",
            "result": "pass",
        })


def postprocess(
    input_docx: Path,
    output_docx: Path,
    profile: str,
    cover_docx: Path | None = None,
    cover_fields: dict[str, str] | None = None,
    abstract_fields: dict[str, str] | None = None,
) -> Path:
    """Add cover, headers, footers, TOC, section breaks, and table formatting."""
    doc = Document(str(input_docx))
    try:
        _set_style_line_spacing_exact(doc.styles["Normal"], "400", "exact")
        _set_style_east_asia_font(doc.styles["Normal"], _SONG)
    except KeyError:
        pass
    fields = {key: value for key, value in (cover_fields or {}).items() if value}
    explicit_title = fields.get("title", "")
    title = explicit_title or _find_title(doc)
    fields.setdefault("title", title)
    header_text = _profile_header_text(profile)
    toc_title = _profile_toc_title(profile)
    margins = PROFILE_MARGINS.get(profile, {"top": 1440, "bottom": 1440, "left": 1440, "right": 1440, "header": 720, "footer": 720})
    paper_size = PROFILE_PAPER_SIZES.get(profile, {"w": 11906, "h": 16838})

    if explicit_title:
        _remove_generated_title_paragraph(doc, explicit_title)

    insert_at = 0

    # 0. Insert selected school cover before generated front matter/body.
    if cover_docx is not None:
        insert_at = _insert_cover(doc, cover_docx, fields, margins=margins, paper_size=paper_size)

    # 0b. Insert required Chinese/English abstract pages after the cover.
    if abstract_fields:
        insert_at = _insert_abstract_pages(doc, insert_at, abstract_fields)

    # 0c. Remove any static/hand-written TOC entries left by Pandoc, then
    #     insert the auto TOC field after cover + abstract, before body text.
    _remove_existing_toc_entries(doc)
    _add_toc(doc, toc_title, insert_at=insert_at)

    # 0.5 Add chapter/section numbering (第一章, 1.1, 1.1.1, ...) before section breaks
    _auto_number_headings(doc)

    # 1. Add section breaks and chapter page breaks
    _add_section_breaks_and_page_breaks(doc, margins=margins, paper_size=paper_size)

    # 2. Format all tables as three-line tables (三线表)
    _format_tables_three_line(doc)

    # 2.2 Format table captions and cell content (font, spacing, indent)
    _format_table_captions_and_content(doc)

    # 2.5 Format figure captions (centered, bold, 10.5pt, spacing)
    _format_figure_captions(doc)

    # 2.8 Format equation numbers (right-aligned, sequential numbering)
    _format_equation_numbers(doc)

    # 2.9 Format headings (一级/二级/三级/四级) and body text per profile spec
    _format_headings_and_body(doc, profile)

    # 3. Set headers and footers on each section
    _set_section_headers_footers(doc, header_text, title)

    # 4.5 Format TOC entries (宋体 小四, left-aligned, no indent)
    _format_toc_entries(doc)

    # 4.6 Move Bibliography entries before 致谢 (Pandoc puts refs at document end)
    _reorder_bibliography_before_thanks(doc)

    # 4.7 Format Bibliography entries (宋体 小四)
    _format_bibliography_entries(doc)

    # Save
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))

    # Post-save: merge cover media, relationships, and content types into the package.
    if cover_docx is not None:
        _merge_cover_package_resources(output_docx, cover_docx)

    return output_docx


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Pandoc-generated DOCX.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    parser.add_argument(
        "--profile",
        required=True,
        choices=sorted(PROFILE_HEADER_TEXT),
        help="Chinese thesis formatting profile.",
    )
    parser.add_argument("--cover-docx", help="Optional DOCX cover template to insert.")
    args = parser.parse_args()

    try:
        result = postprocess(
            input_docx=Path(args.input),
            output_docx=Path(args.output),
            profile=args.profile,
            cover_docx=Path(args.cover_docx) if args.cover_docx else None,
        )
        print(result)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
