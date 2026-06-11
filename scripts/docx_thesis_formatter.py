#!/usr/bin/env python3
"""DOCX thesis final formatter for Chinese university profiles.

First-generation safe formatter:
- reads an input DOCX;
- writes a separate formatted DOCX copy;
- emits a Markdown change report;
- never overwrites the input file.
"""
from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - exercised only on missing deps
    print(
        "[ARS-DOCX-FORMATTER ERROR: python-docx is required. "
        "Install it with `python3 -m pip install python-docx`.]",
        file=sys.stderr,
    )
    raise SystemExit(2) from exc


@dataclass(frozen=True)
class StyleSpec:
    east_asia_font: str
    latin_font: str
    size_pt: float
    bold: bool = False
    alignment: int | None = None
    first_line_indent_cm: float | None = None
    line_spacing_pt: float | None = None


@dataclass(frozen=True)
class ProfileSpec:
    key: str
    label: str
    page_width_cm: float
    page_height_cm: float
    margin_top_cm: float
    margin_bottom_cm: float
    margin_left_cm: float
    margin_right_cm: float
    body: StyleSpec
    headings: dict[int, StyleSpec]
    caption: StyleSpec
    header_left: str | None
    manual_notes: tuple[str, ...]


PROFILES: dict[str, ProfileSpec] = {
    "mainland-fallback": ProfileSpec(
        key="mainland-fallback",
        label="Mainland China University Thesis Fallback",
        page_width_cm=21.0,
        page_height_cm=29.7,
        margin_top_cm=2.54,
        margin_bottom_cm=2.54,
        margin_left_cm=2.5,
        margin_right_cm=2.5,
        body=StyleSpec("宋体", "Times New Roman", 12, first_line_indent_cm=0.74, line_spacing_pt=20),
        headings={
            1: StyleSpec("黑体", "Times New Roman", 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER),
            2: StyleSpec("黑体", "Times New Roman", 15, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            3: StyleSpec("黑体", "Times New Roman", 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            4: StyleSpec("黑体", "Times New Roman", 12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
        },
        caption=StyleSpec("黑体", "Times New Roman", 10.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER),
        header_left=None,
        manual_notes=(
            "Official school or college templates outrank this fallback profile.",
            "Open the output in Word/WPS/LibreOffice and refresh the table of contents and pagination fields.",
        ),
    ),
    "guangxi-undergrad": ProfileSpec(
        key="guangxi-undergrad",
        label="Guangxi University Undergraduate Thesis/Design",
        page_width_cm=21.0,
        page_height_cm=29.7,
        margin_top_cm=2.54,
        margin_bottom_cm=2.54,
        margin_left_cm=2.2,
        margin_right_cm=2.2,
        body=StyleSpec("宋体", "Times New Roman", 12, first_line_indent_cm=0.74, line_spacing_pt=20),
        headings={
            1: StyleSpec("黑体", "Times New Roman", 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER),
            2: StyleSpec("黑体", "Times New Roman", 15, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            3: StyleSpec("黑体", "Times New Roman", 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            4: StyleSpec("黑体", "Times New Roman", 12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
        },
        caption=StyleSpec("黑体", "Times New Roman", 10.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER),
        header_left="广西大学本科毕业论文/毕业设计",
        manual_notes=(
            "Use the official Guangxi University cover and signature pages when available.",
            "Open the output in Word/WPS/LibreOffice and refresh the table of contents and pagination fields.",
            "Check double-sided printing and chapter starts manually before submission/archive.",
        ),
    ),
    "sichuan-grad": ProfileSpec(
        key="sichuan-grad",
        label="Sichuan University Master/Doctoral Dissertation",
        page_width_cm=18.4,
        page_height_cm=26.0,
        margin_top_cm=2.54,
        margin_bottom_cm=2.54,
        margin_left_cm=2.5,
        margin_right_cm=2.5,
        body=StyleSpec("宋体", "Times New Roman", 12, first_line_indent_cm=0.74, line_spacing_pt=20),
        headings={
            1: StyleSpec("黑体", "Times New Roman", 15, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            2: StyleSpec("黑体", "Times New Roman", 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            3: StyleSpec("黑体", "Times New Roman", 12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT),
            4: StyleSpec("楷体", "Times New Roman", 12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT),
        },
        caption=StyleSpec("黑体", "Times New Roman", 10.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER),
        header_left=None,
        manual_notes=(
            "Use the official Sichuan University cover, originality declaration, and authorization statement.",
            "16 kai page dimensions are applied as 18.4 cm x 26.0 cm; verify against the current official template.",
            "Open the output in Word/WPS/LibreOffice and refresh the table of contents and pagination fields.",
        ),
    ),
}


def _set_rfonts(element, *, east_asia: str, latin: str) -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def _apply_style_spec(style, spec: StyleSpec) -> None:
    style.font.name = spec.latin_font
    style.font.size = Pt(spec.size_pt)
    style.font.bold = spec.bold
    _set_rfonts(style.element, east_asia=spec.east_asia_font, latin=spec.latin_font)

    paragraph_format = style.paragraph_format
    if spec.alignment is not None:
        paragraph_format.alignment = spec.alignment
    if spec.first_line_indent_cm is not None:
        paragraph_format.first_line_indent = Cm(spec.first_line_indent_cm)
    if spec.line_spacing_pt is not None:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(spec.line_spacing_pt)


def _ensure_caption_style(document) -> str:
    style_name = "ARS Thesis Caption"
    try:
        document.styles[style_name]
    except KeyError:
        document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    return style_name


def _apply_page_setup(document, profile: ProfileSpec) -> list[str]:
    applied: list[str] = []
    for idx, section in enumerate(document.sections, start=1):
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(profile.page_width_cm)
        section.page_height = Cm(profile.page_height_cm)
        section.top_margin = Cm(profile.margin_top_cm)
        section.bottom_margin = Cm(profile.margin_bottom_cm)
        section.left_margin = Cm(profile.margin_left_cm)
        section.right_margin = Cm(profile.margin_right_cm)
        applied.append(
            f"Section {idx}: page {profile.page_width_cm:g} cm x {profile.page_height_cm:g} cm; "
            f"margins T/B/L/R {profile.margin_top_cm:g}/{profile.margin_bottom_cm:g}/"
            f"{profile.margin_left_cm:g}/{profile.margin_right_cm:g} cm"
        )
    return applied


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def _apply_headers_and_footers(document, profile: ProfileSpec, title: str | None) -> list[str]:
    applied: list[str] = []
    for idx, section in enumerate(document.sections, start=1):
        header_text_parts = []
        if profile.header_left:
            header_text_parts.append(profile.header_left)
        if title:
            header_text_parts.append(title)
        if header_text_parts:
            paragraph = section.header.paragraphs[0]
            _clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run("    ".join(header_text_parts))
            applied.append(f"Section {idx}: header set to {' / '.join(header_text_parts)}")

        footer = section.footer.paragraphs[0]
        _clear_paragraph(footer)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(footer)
        applied.append(f"Section {idx}: centered PAGE field added to footer")
    return applied


def _paragraph_kind(paragraph) -> str:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if style_name.startswith("Heading"):
        return style_name
    if text.startswith(("图", "表", "Fig.", "Figure", "Table")):
        return "caption"
    return "body"


def _apply_paragraph_styles(document, profile: ProfileSpec) -> dict[str, int]:
    counts = {"body": 0, "heading": 0, "caption": 0}
    caption_style_name = _ensure_caption_style(document)
    _apply_style_spec(document.styles[caption_style_name], profile.caption)

    _apply_style_spec(document.styles["Normal"], profile.body)
    for level, spec in profile.headings.items():
        style_name = f"Heading {level}"
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        _apply_style_spec(style, spec)

    for paragraph in document.paragraphs:
        kind = _paragraph_kind(paragraph)
        if kind == "caption":
            paragraph.style = document.styles[caption_style_name]
            counts["caption"] += 1
        elif kind.startswith("Heading"):
            counts["heading"] += 1
        elif paragraph.text.strip():
            paragraph.style = document.styles["Normal"]
            counts["body"] += 1
    return counts


def _has_toc_like_content(document) -> bool:
    if "TOC" in document._element.xml:
        return True
    return any(p.text.strip() in {"目录", "目  录", "Table of Contents"} for p in document.paragraphs[:80])


def _write_report(
    *,
    report_path: Path,
    input_path: Path,
    output_path: Path,
    profile: ProfileSpec,
    applied: Iterable[str],
    counts: dict[str, int],
    has_toc: bool,
    title: str | None,
) -> None:
    applied_lines = "\n".join(f"- {item}" for item in applied)
    manual_lines = "\n".join(f"- {item}" for item in profile.manual_notes)
    toc_status = (
        "TOC-like content or Word TOC field detected; refresh it in Word/WPS/LibreOffice."
        if has_toc
        else "No TOC-like content detected in the first 80 paragraphs; generate or insert the TOC in Word/WPS/LibreOffice after verifying headings."
    )
    body = f"""# DOCX Thesis Formatting Report

| Item | Value |
|------|-------|
| Input | `{input_path}` |
| Output | `{output_path}` |
| Profile | `{profile.key}` — {profile.label} |
| Thesis title | {title or "not provided"} |
| Body paragraphs styled | {counts.get("body", 0)} |
| Heading paragraphs observed | {counts.get("heading", 0)} |
| Caption paragraphs styled | {counts.get("caption", 0)} |

## Applied Changes

{applied_lines}

## Table Of Contents And Pagination

- {toc_status}
- The footer contains a Word `PAGE` field. Open the document in Word/WPS/LibreOffice and update fields before final submission.

## Manual Verification Required

{manual_lines}

## Safety Notes

- The original input DOCX was not overwritten.
- This first-generation formatter does not generate school cover pages, seals, signatures, or final declarations from memory.
- Visual fidelity still requires a later render-and-compare workflow.
"""
    report_path.write_text(body, encoding="utf-8")


def format_docx(
    *,
    input_path: Path,
    output_path: Path,
    profile_key: str,
    title: str | None,
    report_path: Path,
    force_overwrite_output: bool = False,
) -> int:
    input_path = input_path.resolve()
    output_path = output_path.resolve()
    report_path = report_path.resolve()

    if profile_key not in PROFILES:
        raise ValueError(f"unknown profile: {profile_key}")
    if not input_path.exists():
        print(f"[ARS-DOCX-FORMATTER ERROR: input not found: {input_path}]", file=sys.stderr)
        return 2
    if input_path.suffix.lower() != ".docx":
        print("[ARS-DOCX-FORMATTER ERROR: input must be a .docx file]", file=sys.stderr)
        return 2
    if input_path == output_path:
        print("[ARS-DOCX-FORMATTER ERROR: output path must differ from input path]", file=sys.stderr)
        return 2
    if output_path.exists() and not force_overwrite_output:
        print(
            f"[ARS-DOCX-FORMATTER ERROR: output already exists: {output_path}. "
            "Use --force-overwrite-output to replace it.]",
            file=sys.stderr,
        )
        return 2

    output_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    profile = PROFILES[profile_key]
    document = Document(str(input_path))
    applied = []
    applied.extend(_apply_page_setup(document, profile))
    counts = _apply_paragraph_styles(document, profile)
    applied.append("Normal style and Heading 1-4 style definitions updated")
    applied.append("Figure/table caption-like paragraphs styled with ARS Thesis Caption")
    applied.extend(_apply_headers_and_footers(document, profile, title))
    has_toc = _has_toc_like_content(document)

    document.save(str(output_path))
    _write_report(
        report_path=report_path,
        input_path=input_path,
        output_path=output_path,
        profile=profile,
        applied=applied,
        counts=counts,
        has_toc=has_toc,
        title=title,
    )
    print(f"[ars-docx-thesis-formatter] wrote {output_path}")
    print(f"[ars-docx-thesis-formatter] wrote {report_path}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="docx_thesis_formatter",
        description="Create a formatted Chinese university thesis DOCX copy and Markdown report.",
    )
    parser.add_argument("input", help="Input thesis .docx path.")
    parser.add_argument(
        "--profile",
        choices=sorted(PROFILES),
        default="mainland-fallback",
        help="Built-in thesis formatting profile.",
    )
    parser.add_argument("--output", required=True, help="Output .docx path. Must differ from input.")
    parser.add_argument("--title", help="Optional thesis title for supported headers.")
    parser.add_argument("--report", help="Markdown report path. Defaults to <output>.format-report.md.")
    parser.add_argument(
        "--force-overwrite-output",
        action="store_true",
        help="Allow replacing an existing output file. The input file is still never overwritten.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    output_path = Path(args.output)
    report_path = Path(args.report) if args.report else output_path.with_suffix(".format-report.md")
    try:
        return format_docx(
            input_path=Path(args.input),
            output_path=output_path,
            profile_key=args.profile,
            title=args.title,
            report_path=report_path,
            force_overwrite_output=args.force_overwrite_output,
        )
    except ValueError as exc:
        print(f"[ARS-DOCX-FORMATTER ERROR: {exc}]", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
