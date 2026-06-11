from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_thesis_formatter import main


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("这是正文段落，用于测试排版脚本。")
    doc.add_paragraph("图1-1 测试图题")
    doc.save(path)


def test_formatter_writes_copy_and_report(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    report = tmp_path / "report.md"
    _make_docx(source)

    rc = main(
        [
            str(source),
            "--profile",
            "guangxi-undergrad",
            "--output",
            str(output),
            "--title",
            "测试论文题目",
            "--report",
            str(report),
        ]
    )

    assert rc == 0
    assert output.exists()
    assert report.exists()
    report_text = report.read_text(encoding="utf-8")
    assert "Guangxi University Undergraduate Thesis/Design" in report_text
    assert "测试论文题目" in report_text
    assert "refresh" in report_text


def test_formatter_refuses_to_overwrite_input(tmp_path):
    source = tmp_path / "source.docx"
    _make_docx(source)

    rc = main([str(source), "--output", str(source)])

    assert rc == 2


def test_formatter_refuses_existing_output_without_force(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    _make_docx(source)
    _make_docx(output)

    rc = main([str(source), "--output", str(output)])

    assert rc == 2
