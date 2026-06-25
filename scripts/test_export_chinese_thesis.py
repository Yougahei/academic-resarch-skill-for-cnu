from __future__ import annotations

import re
from pathlib import Path
import zipfile

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

import scripts.build_chinese_reference_docx as build_ref
import scripts.export_chinese_thesis as export
from scripts.chinese_fonts import resolve_cjk_font
from scripts.postprocess_chinese_thesis_docx import (
    _add_section_breaks_and_page_breaks,
    _format_equation_numbers,
    _format_figure_captions,
    _format_table_captions_and_content,
    _is_chapter_heading,
    _is_figure_table,
    _paragraph_has_content,
    _paragraph_text,
    postprocess,
)


def test_profile_assets_exist() -> None:
    for profile in export.PROFILES.values():
        assert profile.tex_template.exists()
        assert profile.reference_docx.exists()
        if profile.cover_docx:
            assert profile.cover_docx.exists()


def test_docx_reference_files_are_valid_zip_archives() -> None:
    for profile in export.PROFILES.values():
        with zipfile.ZipFile(profile.reference_docx) as archive:
            names = set(archive.namelist())
            assert "[Content_Types].xml" in names
            assert "word/styles.xml" in names
            styles = archive.read("word/styles.xml").decode("utf-8")
            assert "Times New Roman" in styles
            assert "宋体" in styles
            assert "黑体" in styles


def test_docx_cover_files_are_valid_zip_archives() -> None:
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None
    with zipfile.ZipFile(profile.cover_docx) as archive:
        names = set(archive.namelist())
        assert "[Content_Types].xml" in names
        assert "word/document.xml" in names
    doc = Document(str(profile.cover_docx))
    text = _all_docx_text(doc)
    assert "本科毕业论文" in text
    assert "摘要" not in text


def test_build_pandoc_args_docx_uses_reference_docx(tmp_path: Path) -> None:
    profile = export.PROFILES["guangxi-undergrad"]
    args = export.build_pandoc_args(
        pandoc="pandoc",
        input_path=tmp_path / "paper.md",
        output_path=tmp_path / "paper.docx",
        output_format="docx",
        profile=profile,
    )
    assert "--reference-doc" in args
    assert str(profile.reference_docx) in args
    assert "--template" not in args


def test_extract_cover_fields_uses_frontmatter_and_h1_fallback(tmp_path: Path) -> None:
    paper = tmp_path / "paper.md"
    paper.write_text(
        "\n".join(
            [
                "---",
                "college: 管理学院",
                "major: 信息管理",
                "class-name: 信管 2201",
                "student-id: 20260001",
                "author: 张三",
                "advisor: 李老师",
                "date: 二〇二六年六月",
                "---",
                "# LLM 时代企业知识基础设施研究",
                "",
                "## 第1章 绪论",
            ]
        ),
        encoding="utf-8",
    )

    fields = export.extract_cover_fields(paper, export.PROFILES["guangxi-undergrad"])

    assert fields["title"] == "LLM 时代企业知识基础设施研究"
    assert fields["college"] == "管理学院"
    assert fields["student-id"] == "20260001"
    assert fields["paper-type"] == "本科毕业论文"


def test_prepare_markdown_uses_frontmatter_abstracts_and_removes_h1(tmp_path: Path) -> None:
    paper = tmp_path / "paper.md"
    paper.write_text(
        "\n".join(
            [
                "---",
                "title: 元数据标题",
                "title-en: Metadata Title",
                "abstract-zh: 中文摘要内容。",
                "keywords-zh: 关键词一；关键词二",
                "abstract-en: English abstract content.",
                "keywords-en: keyword one; keyword two",
                "---",
                "# 正文里不应重复的标题",
                "",
                "## 第1章 绪论",
                "正文内容",
            ]
        ),
        encoding="utf-8",
    )

    prepared = export.prepare_markdown_for_export(paper, tmp_path)
    normalized = prepared.path.read_text(encoding="utf-8")

    assert prepared.metadata["title"] == "元数据标题"
    assert prepared.metadata["abstract-zh"] == "中文摘要内容。"
    assert "# 正文里不应重复的标题" not in normalized
    assert "## 第1章 绪论" in normalized


def test_prepare_markdown_extracts_inline_abstract_blocks(tmp_path: Path) -> None:
    paper = tmp_path / "paper.md"
    paper.write_text(
        "\n".join(
            [
                "# 论文题目",
                "",
                "## 摘要",
                "中文摘要第一段。",
                "",
                "**关键词**：本体；大模型",
                "",
                "## ABSTRACT",
                "English abstract paragraph.",
                "",
                "Keywords: ontology; LLM",
                "",
                "## 第1章 绪论",
                "正文内容",
            ]
        ),
        encoding="utf-8",
    )

    prepared = export.prepare_markdown_for_export(paper, tmp_path)
    normalized = prepared.path.read_text(encoding="utf-8")

    assert prepared.metadata["title"] == "论文题目"
    assert prepared.metadata["abstract-zh"] == "中文摘要第一段。"
    assert prepared.metadata["keywords-zh"] == "本体；大模型"
    assert prepared.metadata["abstract-en"] == "English abstract paragraph."
    assert prepared.metadata["keywords-en"] == "ontology; LLM"
    assert "## 摘要" not in normalized
    assert "## ABSTRACT" not in normalized
    assert "# 论文题目" not in normalized
    assert "## 第1章 绪论" in normalized


def test_prepare_markdown_block_scalar_abstracts_without_pyyaml(tmp_path: Path, monkeypatch) -> None:
    import sys

    monkeypatch.setitem(sys.modules, "yaml", None)
    paper = tmp_path / "paper.md"
    paper.write_text(
        "\n".join(
            [
                "---",
                "title: 块标量摘要测试",
                "abstract-zh: |",
                "  中文摘要第一段。",
                "",
                "  中文摘要第二段。",
                "keywords-zh: 关键词一；关键词二",
                "abstract-en: >",
                "  English abstract first paragraph.",
                "",
                "  English abstract second paragraph.",
                "keywords-en: keyword one; keyword two",
                "---",
                "## 第1章 绪论",
                "正文内容",
            ]
        ),
        encoding="utf-8",
    )

    prepared = export.prepare_markdown_for_export(paper, tmp_path)

    assert prepared.metadata["abstract-zh"] != "|"
    assert "中文摘要第一段。" in prepared.metadata["abstract-zh"]
    assert "中文摘要第二段。" in prepared.metadata["abstract-zh"]
    assert "English abstract first paragraph." in prepared.metadata["abstract-en"]
    assert "English abstract second paragraph." in prepared.metadata["abstract-en"]
    assert prepared.metadata["keywords-zh"] == "关键词一；关键词二"
    assert prepared.metadata["keywords-en"] == "keyword one; keyword two"


def test_validate_abstract_metadata_requires_bilingual_abstracts() -> None:
    try:
        export.validate_abstract_metadata({"abstract-zh": "中文摘要"})
    except ValueError as exc:
        message = str(exc)
    else:
        raise AssertionError("Expected missing abstract metadata to fail")

    assert "abstract-en" in message
    assert "keywords-zh" in message
    assert "Run the abstract generation flow first" in message


def test_build_pandoc_args_pdf_uses_tex_template(tmp_path: Path) -> None:
    profile = export.PROFILES["sichuan-grad"]
    args = export.build_pandoc_args(
        pandoc="pandoc",
        input_path=tmp_path / "paper.md",
        output_path=tmp_path / "paper.pdf",
        output_format="pdf",
        profile=profile,
    )
    assert "--template" in args
    assert str(profile.tex_template) in args
    assert "--pdf-engine" in args
    assert "xelatex" in args


def test_build_pandoc_args_pdf_injects_platform_fonts(tmp_path: Path, monkeypatch) -> None:
    from scripts import chinese_fonts

    monkeypatch.setattr(chinese_fonts, "detect_font_platform", lambda: "windows")
    profile = export.PROFILES["guangxi-undergrad"]
    args = export.build_pandoc_args(
        pandoc="pandoc",
        input_path=tmp_path / "paper.md",
        output_path=tmp_path / "paper.pdf",
        output_format="pdf",
        profile=profile,
    )
    assert "-V" in args
    expected = f"cjkfont={chinese_fonts.latex_font_vars('windows')['cjkfont']}"
    assert expected in args


def test_build_pandoc_args_latex_injects_platform_fonts(tmp_path: Path, monkeypatch) -> None:
    from scripts import chinese_fonts

    monkeypatch.setattr(chinese_fonts, "detect_font_platform", lambda: "windows")
    profile = export.PROFILES["guangxi-undergrad"]
    args = export.build_pandoc_args(
        pandoc="pandoc",
        input_path=tmp_path / "paper.md",
        output_path=tmp_path / "paper.tex",
        output_format="latex",
        profile=profile,
    )
    assert "-V" in args
    expected = f"cjkfont={chinese_fonts.latex_font_vars('windows')['cjkfont']}"
    assert expected in args


def test_build_pandoc_args_docx_does_not_inject_fonts(tmp_path: Path) -> None:
    profile = export.PROFILES["guangxi-undergrad"]
    args = export.build_pandoc_args(
        pandoc="pandoc",
        input_path=tmp_path / "paper.md",
        output_path=tmp_path / "paper.docx",
        output_format="docx",
        profile=profile,
    )
    assert "-V" not in args


def _write_minimal_generated_docx(path: Path, include_toc: bool = True) -> None:
    doc = Document()
    if include_toc:
        doc.add_paragraph("目  录")
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("正文内容")
    doc.save(path)


def _write_generated_docx_with_pandoc_title(path: Path, title: str) -> None:
    doc = Document()
    doc.add_paragraph(title)
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("正文内容")
    doc.save(path)


def _all_docx_text(doc: Document) -> str:
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _paragraph_spacing_attr(para, attr: str) -> str | None:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        return None
    return spacing.get(qn(f"w:{attr}"))


def _paragraph_first_line_chars(para) -> str | None:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None
    return ind.get(qn("w:firstLineChars"))


def _run_has_bold(run) -> bool:
    rPr = run._element.find(qn("w:rPr"))
    return rPr is not None and rPr.find(qn("w:b")) is not None


def test_postprocess_inserts_guangxi_cover_before_toc(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={
            "title": "LLM 时代企业知识基础设施研究",
            "college": "管理学院",
            "major": "信息管理",
            "class-name": "信管 2201",
            "student-id": "20260001",
            "author": "张三",
            "advisor": "李老师",
            "date": "二〇二六年六月",
        },
    )

    doc = Document(str(output_docx))
    paragraph_texts = [para.text for para in doc.paragraphs]
    assert paragraph_texts.index("本科毕业论文") < paragraph_texts.index("目  录")

    all_text = _all_docx_text(doc)
    assert "LLM 时代企业知识基础设施研究" in all_text
    assert "管理学院" in all_text
    assert "20260001" in all_text
    assert "二〇二六年六月" in all_text


def test_postprocess_inserts_abstract_pages_between_cover_and_toc(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_generated_docx_with_pandoc_title(input_docx, "封面标题")
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": "封面标题"},
        abstract_fields={
            "title-en": "the impact of AI on learning engagement",
            "abstract-zh": "中文摘要内容。",
            "keywords-zh": "本体；大模型。",
            "abstract-en": "English abstract content.",
            "keywords-en": "ontology; LLM",
        },
    )

    doc = Document(str(output_docx))
    paragraph_texts = [para.text for para in doc.paragraphs]
    assert paragraph_texts.index("本科毕业论文") < paragraph_texts.index("摘　要")
    assert paragraph_texts.index("摘　要") < paragraph_texts.index("ABSTRACT")
    assert paragraph_texts.index("ABSTRACT") < paragraph_texts.index("目  录")
    assert paragraph_texts.index("目  录") < paragraph_texts.index("第1章 绪论")
    # The title now appears once as a standalone cover paragraph (Issue #79
    # moved it out of the cover table into an independent paragraph).  It
    # must appear exactly once and before the abstract heading.
    title_count = paragraph_texts.count("封面标题")
    assert title_count == 1, f"title should appear once as cover paragraph, got {title_count}"
    assert paragraph_texts.index("封面标题") < paragraph_texts.index("摘　要")
    all_text = _all_docx_text(doc)
    assert "中文摘要内容。" in all_text
    assert "English abstract content." in all_text
    assert "本体；大模型。" not in all_text
    assert "本体；大模型" in all_text
    assert "The Impact Of AI On Learning Engagement" in all_text
    assert "Ontology; LLM" in all_text

    zh_heading_idx = paragraph_texts.index("摘　要")
    assert paragraph_texts[zh_heading_idx - 1] == ""
    assert paragraph_texts[zh_heading_idx + 1] == ""
    zh_heading = doc.paragraphs[zh_heading_idx]
    assert zh_heading.alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER
    assert _paragraph_spacing_attr(zh_heading, "line") == "400"
    assert zh_heading.runs[0].text == "摘　要"
    assert _run_has_bold(zh_heading.runs[0])

    keyword_idx = paragraph_texts.index("关键词：本体；大模型")
    assert paragraph_texts[keyword_idx - 1] == ""
    keyword_para = doc.paragraphs[keyword_idx]
    assert _paragraph_first_line_chars(keyword_para) == "200"
    assert [run.text for run in keyword_para.runs] == ["关键词：", "本体；大模型"]
    assert _run_has_bold(keyword_para.runs[0])
    assert not _run_has_bold(keyword_para.runs[1])

    en_title_idx = paragraph_texts.index("The Impact Of AI On Learning Engagement")
    assert paragraph_texts[en_title_idx + 1] == ""
    abstract_idx = paragraph_texts.index("ABSTRACT")
    assert paragraph_texts[abstract_idx + 1] == ""
    keywords_idx = paragraph_texts.index("Keywords: Ontology; LLM")
    assert paragraph_texts[keywords_idx - 1] == ""
    keywords_para = doc.paragraphs[keywords_idx]
    assert [run.text for run in keywords_para.runs] == ["Keywords: ", "Ontology; LLM"]
    assert _run_has_bold(keywords_para.runs[0])
    assert not _run_has_bold(keywords_para.runs[1])


def test_postprocess_adds_missing_toc_after_cover(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx, include_toc=False)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": "真实导出无目录场景"},
    )

    doc = Document(str(output_docx))
    paragraph_texts = [para.text for para in doc.paragraphs]
    assert paragraph_texts.index("本科毕业论文") < paragraph_texts.index("目  录")
    assert paragraph_texts.index("目  录") < paragraph_texts.index("第1章 绪论")


def test_postprocess_keeps_missing_cover_fields_blank(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": "只填标题"},
    )

    doc = Document(str(output_docx))
    # Issue #79 removed the 课题名称 table, so the field table (学院/专业/...)
    # is now the first table in the document.
    cover_info = doc.tables[0]
    values_by_label = {row.cells[0].text.strip(): row.cells[1].text for row in cover_info.rows}
    assert values_by_label["专业："] == ""
    assert values_by_label["班级："] == ""
    assert values_by_label["学号："] == ""
    assert values_by_label["姓名："] == ""
    assert values_by_label["指导老师："] == ""


def test_postprocess_without_cover_does_not_insert_guangxi_cover(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile="sichuan-grad",
    )

    doc = Document(str(output_docx))
    assert "本科毕业论文" not in _all_docx_text(doc)


def test_postprocess_paragraph_order_cover_toc_body(tmp_path: Path) -> None:
    """Regression: cover title must appear before TOC, and TOC before first chapter.

    Verifies the structural (OOXML body-element) order, not just paragraph text
    search, so that reordering bugs are caught even when the expected text is
    present somewhere in the document.
    """
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_generated_docx_with_pandoc_title(input_docx, "测试论文题目")
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": "测试论文题目"},
        abstract_fields={
            "title-en": "test title",
            "abstract-zh": "中文摘要。",
            "keywords-zh": "关键词A；关键词B",
            "abstract-en": "English abstract.",
            "keywords-en": "keywordA; keywordB",
        },
    )

    doc = Document(str(output_docx))
    body_children = list(doc.element.body)

    # Locate landmark elements by their text content
    def _elem_index_containing(
        children, text_fragment: str, skip_tag: str | None = None
    ) -> int:
        for i, elem in enumerate(children):
            if skip_tag and elem.tag == skip_tag:
                continue
            all_t = "".join(
                t.text or "" for t in elem.findall(".//" + qn("w:t"))
            )
            if text_fragment in all_t:
                return i
        return -1

    # Exclude sectPr elements (section properties live at body level)
    w_p = qn("w:p")
    cover_idx = _elem_index_containing(body_children, "本科毕业论文", skip_tag=None)
    toc_heading_idx = _elem_index_containing(body_children, "目  录", skip_tag=None)
    toc_field_idx = _elem_index_containing(body_children, "请在Word中右键更新目录", skip_tag=None)
    first_chap_idx = _elem_index_containing(body_children, "第1章", skip_tag=None)

    assert cover_idx >= 0, "cover title must be in body"
    assert toc_heading_idx >= 0, "TOC heading must be in body"
    assert toc_field_idx >= 0, "TOC field must be in body"
    assert first_chap_idx >= 0, "first chapter must be in body"

    assert cover_idx < toc_heading_idx, (
        f"cover ({cover_idx}) must appear before TOC heading ({toc_heading_idx})"
    )
    assert toc_heading_idx < toc_field_idx, (
        f"TOC heading ({toc_heading_idx}) must appear before TOC field ({toc_field_idx})"
    )
    assert toc_field_idx < first_chap_idx, (
        f"TOC field ({toc_field_idx}) must appear before first chapter ({first_chap_idx})"
    )


def test_build_tool_env_returns_path() -> None:
    env = export.build_tool_env()
    assert "PATH" in env
    assert isinstance(env["PATH"], str)


def test_patch_styles_xml_sets_chinese_fonts() -> None:
    xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/></w:style>
    </w:styles>
    """
    patched = build_ref.patch_styles_xml(xml, build_ref.PROFILES["mainland-fallback"]).decode("utf-8")
    assert "Times New Roman" in patched
    assert "宋体" in patched
    assert "黑体" in patched
    assert "firstLineChars" in patched


# ---------------------------------------------------------------------------
# Profile auto-discovery tests
# ---------------------------------------------------------------------------


def test_discover_profiles_finds_builtin_templates() -> None:
    from scripts.export_chinese_thesis import discover_profiles, PROFILES

    discovered = discover_profiles()
    assert "guangxi-undergrad" in discovered, "should find guangxi reference docx"
    assert "sichuan-grad" in discovered, "should find sichuan reference docx"
    assert "mainland-fallback" in discovered, "should find mainland reference docx"

    # Hardcoded profiles should have proper header/toc values
    p = PROFILES["guangxi-undergrad"]
    assert p.header_text == "广西大学本科毕业论文"
    assert p.toc_title == "目  录"

    p = PROFILES["mainland-fallback"]
    assert p.header_text == "本科毕业论文"
    assert p.toc_title == "目录"


def test_discover_profiles_json_sidecar(tmp_path: Path) -> None:
    import importlib, json
    from scripts.export_chinese_thesis import DOCX_TEMPLATE_DIR, discover_profiles

    # Create a mock reference docx + sidecar JSON
    ref = DOCX_TEMPLATE_DIR / f"test_univ_undergrad_reference.docx"
    meta = DOCX_TEMPLATE_DIR / "test_univ_undergrad.json"
    try:
        ref.write_text("")
        meta.write_text(json.dumps({
            "label": "Test University",
            "header_text": "测试大学本科毕业论文",
            "toc_title": "目  录",
            "school_label": "测试大学",
            "degree_label": "本科",
            "paper_type_label": "本科毕业论文",
        }, ensure_ascii=False))

        # Re-discover (module-level PROFILES already loaded, but discover_profiles re-scans)
        import scripts.export_chinese_thesis as m
        importlib.reload(m)
        p = m.PROFILES.get("test-univ-undergrad")
        assert p is not None, f"test-univ-undergrad should be discovered. Got: {sorted(m.PROFILES.keys())}"
        assert p.header_text == "测试大学本科毕业论文"
        assert p.toc_title == "目  录"
        assert p.label == "Test University"
        assert p.school_label == "测试大学"
    finally:
        ref.unlink(missing_ok=True)
        meta.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Chapter page break tests
# ---------------------------------------------------------------------------

def _has_page_break_before(para) -> bool:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:pageBreakBefore")) is not None


def _has_section_break_before(elem) -> bool:
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:sectPr")) is not None


def test_is_chapter_heading_variants() -> None:
    assert _is_chapter_heading("第1章 绪论")
    assert _is_chapter_heading("第10章 讨论")
    assert _is_chapter_heading("第一章 引言")
    assert _is_chapter_heading("参考文献")
    assert _is_chapter_heading("致  谢")
    assert _is_chapter_heading("致谢")
    assert _is_chapter_heading("附录")
    assert _is_chapter_heading("附录A 调查问卷")


def test_is_chapter_heading_non_matches() -> None:
    assert not _is_chapter_heading("摘要")
    assert not _is_chapter_heading("ABSTRACT")
    assert not _is_chapter_heading("目  录")
    assert not _is_chapter_heading("关键词：本体；大模型")
    assert not _is_chapter_heading("Keywords: ontology; LLM")
    assert not _is_chapter_heading("1.1 背景介绍")
    assert not _is_chapter_heading("")


def test_add_section_breaks_and_page_breaks_first_chapter(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目  录")
    p1 = doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("绪论正文")
    p2 = doc.add_paragraph("第2章 文献综述")
    doc.add_paragraph("文献综述正文")

    _add_section_breaks_and_page_breaks(doc)

    # First chapter should have a section break before it
    first_heading = body_paragraph_by_text(doc, "第1章 绪论")
    assert first_heading is not None
    assert not _has_page_break_before(p1), "first chapter must not get pageBreakBefore"

    # There should be a section break paragraph somewhere before first chapter
    body = doc.element.body
    first_chapter_index = None
    for i, elem in enumerate(body):
        if elem.tag == qn("w:p"):
            text = _paragraph_text(elem)
            if text == "第1章 绪论":
                first_chapter_index = i
                break

    assert first_chapter_index is not None and first_chapter_index > 0
    prev = body[first_chapter_index - 1]
    assert _has_section_break_before(prev), "expected section break before first chapter"


def body_paragraph_by_text(doc: Document, text: str):
    """Find first body paragraph in doc by exact text match."""
    body = doc.element.body
    for elem in body:
        if elem.tag == qn("w:p"):
            if _paragraph_text(elem) == text:
                return elem
    return None


def test_add_section_breaks_and_page_breaks_subsequent_chapters(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目  录")
    p_ch1 = doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("绪论正文")
    p_ch2 = doc.add_paragraph("第2章 文献综述")
    doc.add_paragraph("文献综述正文")
    p_ch3 = doc.add_paragraph("第3章 研究方法")
    doc.add_paragraph("研究方法正文")

    _add_section_breaks_and_page_breaks(doc)

    assert not _has_page_break_before(p_ch1), "first chapter must not get pageBreakBefore"
    assert _has_page_break_before(p_ch2), "chapter 2 must get pageBreakBefore"
    assert _has_page_break_before(p_ch3), "chapter 3 must get pageBreakBefore"


def test_add_section_breaks_and_page_breaks_references_and_appendix(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目  录")
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("绪论正文")
    p_ref = doc.add_paragraph("参考文献")
    p_app = doc.add_paragraph("附录A 调查问卷")

    _add_section_breaks_and_page_breaks(doc)

    assert _has_page_break_before(p_ref), "参考文献 must get pageBreakBefore"
    assert _has_page_break_before(p_app), "附录 must get pageBreakBefore"


def test_add_section_breaks_and_page_breaks_no_chapter(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("正文内容")

    # Must not crash
    _add_section_breaks_and_page_breaks(doc)


def test_add_section_breaks_toc_entries_skipped(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目  录")
    # Add a TOC-style paragraph (simulating what Pandoc generates)
    p_toc = doc.add_paragraph("第1章 绪论")
    p_toc.style = doc.styles["Normal"]
    # Override to TOC style
    pPr = p_toc._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p_toc._element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:pStyle {nsdecls("w")} w:val="toc 1"/>'))

    p_ch1 = doc.add_paragraph("第1章 绪论")  # real heading
    doc.add_paragraph("绪论正文")
    p_ch2 = doc.add_paragraph("第2章 文献综述")
    doc.add_paragraph("文献综述正文")

    _add_section_breaks_and_page_breaks(doc)

    # TOC entry should NOT be detected as chapter heading
    assert not _has_page_break_before(p_ch1), "first chapter still no break"
    assert _has_page_break_before(p_ch2), "chapter 2 must get pageBreakBefore"


# ---------------------------------------------------------------------------
# Table caption and content formatting tests
# ---------------------------------------------------------------------------


def _add_table_with_caption(doc: Document, caption_text: str, rows: int = 3) -> None:
    """Add a TableCaption paragraph + table to a Document."""
    cap_para = doc.add_paragraph(caption_text)
    pPr = cap_para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        cap_para._element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:pStyle {nsdecls("w")} w:val="TableCaption"/>'))

    table = doc.add_table(rows=rows, cols=3)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "列2"
    table.cell(0, 2).text = "列3"
    for r in range(1, rows):
        table.cell(r, 0).text = f"数据{r}-1"
        table.cell(r, 1).text = f"数据{r}-2"


def test_format_table_captions_and_content_basic() -> None:
    doc = Document()
    doc.add_paragraph("正文内容")
    _add_table_with_caption(doc, "表1：实验数据对比")
    doc.add_paragraph("第1章 绪论")

    _format_table_captions_and_content(doc)

    # Find the TableCaption paragraph and check formatting
    caption_elem = None
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "TableCaption":
                caption_elem = elem
                break

    assert caption_elem is not None, "TableCaption paragraph not found"
    cap_pPr = caption_elem.find(qn("w:pPr"))
    jc = cap_pPr.find(qn("w:jc"))
    assert jc is not None and jc.get(qn("w:val")) == "center", "caption should be centered"
    for run in caption_elem.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        assert rPr is not None and rPr.find(qn("w:b")) is not None, "caption should be bold"
        sz = rPr.find(qn("w:sz"))
        assert sz is not None and sz.get(qn("w:val")) == "21", "caption font size should be 21"


def test_format_table_captions_and_content_spacing() -> None:
    doc = Document()
    doc.add_paragraph("正文内容")
    _add_table_with_caption(doc, "表1：实验数据对比")
    doc.add_paragraph("第1章 绪论")

    original_count = len([e for e in doc.element.body if e.tag == qn("w:p")])
    _format_table_captions_and_content(doc)

    new_count = len([e for e in doc.element.body if e.tag == qn("w:p")])
    assert new_count == original_count + 2, "blank lines before caption and after table"


def test_format_table_captions_and_content_cell_format() -> None:
    doc = Document()
    doc.add_paragraph("正文")
    _add_table_with_caption(doc, "表1：测试")

    _format_table_captions_and_content(doc)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._element.find(qn("w:pPr"))
                    if pPr is not None:
                        ind = pPr.find(qn("w:ind"))
                        if ind is not None:
                            flc = ind.get(qn("w:firstLineChars"), "not-set")
                            assert flc == "0", f"firstLineChars should be 0, got {flc}"
                    for run in para.runs:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is not None:
                                ea = rFonts.get(qn("w:eastAsia"))
                                if ea:
                                    assert ea == resolve_cjk_font("宋体"), f"cell font should be {resolve_cjk_font('宋体')}, got {ea}"


def test_format_table_captions_and_content_multiple() -> None:
    doc = Document()
    doc.add_paragraph("正文")
    _add_table_with_caption(doc, "表1：数据1")
    _add_table_with_caption(doc, "表2：数据2")

    _format_table_captions_and_content(doc)

    caption_count = 0
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "TableCaption":
                cap_pPr = elem.find(qn("w:pPr"))
                jc = cap_pPr.find(qn("w:jc"))
                assert jc is not None and jc.get(qn("w:val")) == "center", \
                    f"caption {caption_count + 1} should be centered"
                caption_count += 1
    assert caption_count == 2, "both table captions should be formatted"


def test_format_table_captions_and_content_no_tables() -> None:
    doc = Document()
    doc.add_paragraph("只有正文，没有表格")

    # Must not crash
    _format_table_captions_and_content(doc)


def test_format_table_captions_and_content_no_caption() -> None:
    doc = Document()
    doc.add_paragraph("正文")
    doc.add_table(rows=2, cols=2)  # Table without TableCaption

    # Must not crash
    _format_table_captions_and_content(doc)


# ---------------------------------------------------------------------------
# Figure caption formatting tests
# ---------------------------------------------------------------------------


def _add_style_to_paragraph(para, style_val: str) -> None:
    """Set w:pStyle on a paragraph's XML element."""
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    existing = pPr.find(qn("w:pStyle"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(parse_xml(f'<w:pStyle {nsdecls("w")} w:val="{style_val}"/>'))


def _write_figure_docx(
    path: Path,
    num_figures: int = 1,
    include_caption: bool = True,
    include_toc: bool = False,
) -> None:
    """Create a minimal DOCX with CaptionedFigure + ImageCaption pairs."""
    doc = Document()
    if include_toc:
        doc.add_paragraph("目  录")
    doc.add_paragraph("正文内容")

    for i in range(num_figures):
        fig_para = doc.add_paragraph(f"图{i+1}内容")
        _add_style_to_paragraph(fig_para, "CaptionedFigure")
        if include_caption:
            cap_para = doc.add_paragraph(f"图{i+1} 这是一个图题")
            _add_style_to_paragraph(cap_para, "ImageCaption")

    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("正文内容结束")
    doc.save(path)


def _child_style_vals(doc: Document) -> list[tuple[str, str]]:
    """Return (text, style_val) for each body paragraph."""
    result: list[tuple[str, str]] = []
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        text = _paragraph_text(elem)
        pPr = elem.find(qn("w:pPr"))
        style_val = ""
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style_val = pStyle.get(qn("w:val"), "")
        result.append((text, style_val))
    return result


def _is_centered(elem) -> bool:
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        return False
    return jc.get(qn("w:val")) == "center"


def _all_runs_bold(elem) -> bool:
    for run in elem.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None or rPr.find(qn("w:b")) is None:
            return False
    return bool(elem.findall(qn("w:r")))  # False if no runs


def _run_font_size(elem) -> int | None:
    for run in elem.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                return int(sz.get(qn("w:val"), "0"))
    return None


def test_format_figure_captions_basic(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_figure_docx(path)
    doc = Document(str(path))

    _format_figure_captions(doc)

    # Find the ImageCaption paragraph
    caption_elem = None
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "ImageCaption":
                caption_elem = elem
                break

    assert caption_elem is not None, "ImageCaption paragraph not found"
    assert _is_centered(caption_elem), "caption should be centered"
    assert _all_runs_bold(caption_elem), "caption should be bold"
    assert _run_font_size(caption_elem) == 21, "caption font size should be 21 (10.5pt)"


def test_format_figure_captions_spacing(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_figure_docx(path, num_figures=1)
    doc = Document(str(path))
    original_count = len([e for e in doc.element.body if e.tag == qn("w:p")])

    _format_figure_captions(doc)

    new_count = len([e for e in doc.element.body if e.tag == qn("w:p")])
    assert new_count == original_count + 2, "blank lines before figure and after caption"

    styles = [s for _, s in _child_style_vals(doc)]
    cf_idx = styles.index("CaptionedFigure")
    ic_idx = styles.index("ImageCaption")
    assert cf_idx > 0, "blank paragraph should exist before CaptionedFigure"
    assert ic_idx == cf_idx + 1, "ImageCaption should immediately follow CaptionedFigure"
    # The blank after ImageCaption is at ic_idx + 1 (unstyled paragraph)
    assert ic_idx + 1 < len(styles), "blank paragraph should exist after ImageCaption"


def test_format_figure_captions_multiple(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_figure_docx(path, num_figures=3)
    doc = Document(str(path))

    _format_figure_captions(doc)

    count = 0
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        pPr = elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val")) == "ImageCaption":
                assert _is_centered(elem), f"caption {count + 1} should be centered"
                assert _all_runs_bold(elem), f"caption {count + 1} should be bold"
                count += 1

    assert count == 3, "all 3 figure captions should be found and formatted"


def test_format_figure_captions_no_figures(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("只有正文，没有图")
    doc.add_paragraph("第1章 绪论")

    # Must not crash
    _format_figure_captions(doc)


def test_format_figure_captions_no_caption_following(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_figure_docx(path, include_caption=False)
    doc = Document(str(path))

    # Must not crash when CaptionedFigure has no ImageCaption after it
    _format_figure_captions(doc)


def test_is_figure_table_detection() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Set FigureTable style on the table
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblStyle"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(f'<w:tblStyle {nsdecls("w")} w:val="FigureTable"/>'))

    assert _is_figure_table(table), "FigureTable style table should be detected"


def test_format_tables_three_line_skips_figure_table(tmp_path: Path) -> None:
    from scripts.postprocess_chinese_thesis_docx import _format_tables_three_line

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Set FigureTable style
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblStyle"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(f'<w:tblStyle {nsdecls("w")} w:val="FigureTable"/>'))

    _format_tables_three_line(doc)

    # FigureTable should NOT have three-line table borders
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        top = tblBorders.find(qn("w:top"))
        assert top is None or top.get(qn("w:val")) != "single", "FigureTable should not get three-line borders"


# ---------------------------------------------------------------------------
# Equation formatting tests
# ---------------------------------------------------------------------------

def _add_display_equation(doc: Document, text: str) -> None:
    """Append a display equation paragraph (w:p > m:oMathPara > m:oMath) to a doc."""
    from docx.oxml.ns import nsmap as _nsmap
    body = doc.element.body
    w_ns = _nsmap["w"]
    m_ns = _nsmap["m"]
    ns = f'xmlns:w="{w_ns}" xmlns:m="{m_ns}"'
    para_xml = (
        f'<w:p {ns}>'
        f'  <w:pPr><w:pStyle w:val="BodyText"/></w:pPr>'
        f'  <m:oMathPara>'
        f'    <m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
        f'    <m:oMath>'
        f'      <m:r><m:t>{{}}</m:t></m:r>'
        f'    </m:oMath>'
        f'  </m:oMathPara>'
        f'</w:p>'
    )
    para = parse_xml(para_xml.format(text))
    body.append(para)


def _add_inline_equation_paragraph(doc: Document, before: str, after: str) -> None:
    """Append a paragraph with inline math (text + m:oMath + text)."""
    from docx.oxml.ns import nsmap as _nsmap
    body = doc.element.body
    w_ns = _nsmap["w"]
    m_ns = _nsmap["m"]
    ns = f'xmlns:w="{w_ns}" xmlns:m="{m_ns}"'
    para_xml = (
        f'<w:p {ns}>'
        f'  <w:pPr><w:pStyle w:val="BodyText"/></w:pPr>'
        f'  <w:r><w:t>{before}</w:t></w:r>'
        f'  <m:oMath>'
        f'    <m:r><m:t>a=b</m:t></m:r>'
        f'  </m:oMath>'
        f'  <w:r><w:t>{after}</w:t></w:r>'
        f'</w:p>'
    )
    body.append(parse_xml(para_xml))


def _write_equation_docx(path, num_equations=1, include_inline=False):
    """Create a minimal DOCX with display equations and optional inline math."""
    doc = Document()
    doc.add_paragraph("正文内容")
    for i in range(num_equations):
        _add_display_equation(doc, f"E = mc^{{{i+1}}}")
    if include_inline:
        _add_inline_equation_paragraph(doc, "内联公式", "正文继续")
    doc.add_paragraph("第1章 绪论")
    doc.save(path)


def test_format_equation_numbers_basic(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_equation_docx(path, num_equations=1)
    doc = Document(str(path))

    _format_equation_numbers(doc)

    all_text = "\n".join(p.text for p in doc.paragraphs)
    # "第1章 绪论" appears before equations in the DOCX body (test fixture
    # quirk), so equations get chapter-scoped "(1-1)" rather than global "(1)".
    assert "(1-1)" in all_text, "equation number should be chapter-scoped (1-1)"


def test_format_equation_numbers_sequential(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_equation_docx(path, num_equations=3)
    doc = Document(str(path))

    _format_equation_numbers(doc)

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "(1-1)" in all_text
    assert "(1-2)" in all_text
    assert "(1-3)" in all_text


def test_format_equation_numbers_tab_stop(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_equation_docx(path, num_equations=1)
    doc = Document(str(path))

    _format_equation_numbers(doc)

    # Verify right tab stop was added to the equation paragraph
    found_tab = False
    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        if elem.find(qn("m:oMathPara")) is None:
            continue
        pPr = elem.find(qn("w:pPr"))
        assert pPr is not None, "pPr should exist"
        tabs = pPr.find(qn("w:tabs"))
        assert tabs is not None, "tabs element should exist"
        for tab in tabs:
            if tab.get(qn("w:val")) == "right":
                found_tab = True
                break
    assert found_tab, "right tab stop should be added"


def test_format_equation_numbers_no_equations(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("只有正文，没有公式")
    doc.add_paragraph("第1章 绪论")

    # Must not crash
    _format_equation_numbers(doc)


def test_format_equation_numbers_inline_only(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_equation_docx(path, num_equations=0, include_inline=True)
    doc = Document(str(path))

    _format_equation_numbers(doc)

    # Inline math should NOT be numbered
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert all("(" not in p for p in [p.text for p in doc.paragraphs]), \
        "inline math should not get equation numbers"


def test_format_equation_numbers_spacing(tmp_path: Path) -> None:
    path = tmp_path / "test.docx"
    _write_equation_docx(path, num_equations=1)
    doc = Document(str(path))

    _format_equation_numbers(doc)

    for elem in doc.element.body:
        if elem.tag != qn("w:p"):
            continue
        if elem.find(qn("m:oMathPara")) is None:
            continue
        pPr = elem.find(qn("w:pPr"))
        assert pPr is not None
        spacing = pPr.find(qn("w:spacing"))
        assert spacing is not None, "spacing should be set"
        assert spacing.get(qn("w:line")) == "400", "line spacing should be 400 twips"
        assert spacing.get(qn("w:lineRule")) == "exact", "line rule should be exact"


# ---------------------------------------------------------------------------
# Heading outline level (#75) and Normal style line spacing (#76) tests
# ---------------------------------------------------------------------------


def _outline_level(elem) -> str | None:
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return None
    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    return outline.get(qn("w:val"))


def test_format_headings_sets_outline_level() -> None:
    from scripts.postprocess_chinese_thesis_docx import _format_headings_and_body

    doc = Document()
    h1 = doc.add_paragraph("第1章 绪论")
    _add_style_to_paragraph(h1, "Heading1")
    h2 = doc.add_paragraph("1.1 背景介绍")
    _add_style_to_paragraph(h2, "Heading2")
    h3 = doc.add_paragraph("1.1.1 研究意义")
    _add_style_to_paragraph(h3, "Heading3")

    _format_headings_and_body(doc, "mainland-fallback")

    assert _outline_level(h1._element) == "0"
    assert _outline_level(h2._element) == "1"
    assert _outline_level(h3._element) == "2"


def test_postprocess_sets_normal_style_line_spacing(tmp_path: Path) -> None:
    from scripts.postprocess_chinese_thesis_docx import validate_format

    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["mainland-fallback"]

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
    )

    doc = Document(str(output_docx))
    normal = doc.styles["Normal"]
    pPr = normal.element.find(qn("w:pPr"))
    assert pPr is not None
    spacing = pPr.find(qn("w:spacing"))
    assert spacing is not None
    assert spacing.get(qn("w:line")) == "400"
    assert spacing.get(qn("w:lineRule")) == "exact"

    results = validate_format(output_docx, profile.id)
    spacing_result = next(
        r for r in results if "Normal style" in r["check"]
    )
    assert spacing_result["result"] == "pass", spacing_result


# ---------------------------------------------------------------------------
# Cover deduplication (#73) and compaction (#74) tests
# ---------------------------------------------------------------------------


def test_postprocess_cover_title_not_duplicated(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    title = "LLM 时代企业知识基础设施研究"
    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": title},
    )

    doc = Document(str(output_docx))
    # Issue #79 moved the title out of the cover table into a standalone
    # paragraph.  The title must appear exactly once across all paragraphs
    # and must NOT appear inside any table cell (the 课题名称 table is gone).
    title_paragraph_count = sum(
        1 for para in doc.paragraphs if para.text.strip() == title
    )
    assert title_paragraph_count == 1, (
        f"title should appear once as a standalone paragraph, got {title_paragraph_count}"
    )

    title_in_tables = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2 and row.cells[1].text.strip() == title:
                title_in_tables += 1
    assert title_in_tables == 0, (
        f"title should not appear in any table cell, got {title_in_tables}"
    )


def test_postprocess_cover_compact_no_extra_blank_paragraphs(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={"title": "封面标题"},
    )

    doc = Document(str(output_docx))
    title_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == "本科毕业论文":
            title_idx = idx
            break
    assert title_idx is not None, "cover title paragraph not found"

    # The cover template starts with a logo paragraph (which has a drawing
    # but no text, so para.text is "").  _compact_cover_top now treats
    # paragraphs with drawings as content, so blanks are collapsed between
    # the logo and the "本科毕业论文" title.  At most one blank should
    # appear between any two content paragraphs before the first table.
    # We use _paragraph_has_content() to correctly identify the logo
    # paragraph (which has a drawing) as non-blank.
    blanks_before = sum(
        1 for para in doc.paragraphs[:title_idx]
        if not para.text.strip() and not _paragraph_has_content(para._element)
    )
    assert blanks_before <= 1, f"expected <=1 blank before title, got {blanks_before}"

    # Issue #79 removed the 课题名称 table and replaced it with a standalone
    # title paragraph.  Verify no table contains a 课题名称 label row.
    title_label_rows = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 1 and "课题名称" in row.cells[0].text:
                title_label_rows += 1
    assert title_label_rows == 0, (
        f"课题名称 table should have been removed, found {title_label_rows} label rows"
    )


def test_postprocess_writes_platform_fonts(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _write_minimal_generated_docx(input_docx)
    profile = export.PROFILES["guangxi-undergrad"]
    assert profile.cover_docx is not None

    postprocess(
        input_docx=input_docx,
        output_docx=output_docx,
        profile=profile.id,
        cover_docx=profile.cover_docx,
        cover_fields={
            "title": "跨平台字体测试",
            "advisor": "王老师",
        },
    )

    doc = Document(str(output_docx))
    advisor_ea = None
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = re.sub(r"[\s:：]+", "", row.cells[0].text)
            if label in ("指导老师", "指导教师"):
                for para in row.cells[1].paragraphs:
                    for run in para.runs:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is not None:
                                ea = rFonts.get(qn("w:eastAsia"))
                                if ea:
                                    advisor_ea = ea
    assert advisor_ea is not None, "advisor cell run should have an east-asia font"
    assert advisor_ea == resolve_cjk_font("宋体"), (
        f"advisor font should be {resolve_cjk_font('宋体')}, got {advisor_ea}"
    )
